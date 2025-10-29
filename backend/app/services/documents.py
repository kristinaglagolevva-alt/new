from __future__ import annotations

import json
import logging
import os
import re
from typing import Iterable, Optional, Tuple, List, Any, Collection
from collections import defaultdict
from uuid import uuid4

# env: подхватываем и backend/.env, и корневой .env.local
try:
    from dotenv import load_dotenv  # pip install python-dotenv
    from ..config import BASE_DIR
    from pathlib import Path as _Path
    # backend/.env
    load_dotenv(BASE_DIR / ".env")
    # корневой .env.local (например: /Users/.../Jira Integration Workflow/.env.local)
    load_dotenv(BASE_DIR.parent / ".env.local")
    # корневой .env
    load_dotenv(BASE_DIR.parent / ".env")
    # на всякий случай: текущая рабочая директория
    load_dotenv(_Path.cwd() / ".env.local")
    load_dotenv(_Path.cwd() / ".env")
except Exception:
    pass

# --- GPT integration (опционально) ---
try:
    from openai import OpenAI  # pip install openai>=1.40
except Exception:  # pragma: no cover - optional dependency
    OpenAI = None  # type: ignore

OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
OPENAI_MODEL = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
_openai_client: Optional["OpenAI"] = None
if OPENAI_API_KEY and OpenAI is not None:
    _openai_client = OpenAI(api_key=OPENAI_API_KEY)

from datetime import datetime, date
from calendar import monthrange
from .contracts import ServiceError


def _normalize_role_value(value: Any) -> str | None:
    candidate = getattr(value, "value", value)
    if not isinstance(candidate, str):
        return None
    normalized = candidate.strip().lower()
    if not normalized:
        return None
    return normalized


def _collect_user_roles(primary: Any, extra: Iterable[Any] | None) -> set[str]:
    roles: set[str] = set()
    primary_value = _normalize_role_value(primary)
    if primary_value:
        roles.add(primary_value)
    if extra:
        for item in extra:
            normalized = _normalize_role_value(item)
            if normalized:
                roles.add(normalized)
    return roles


def _v2_work_package_key_from_ids(package_id: str | int | None, performer_id: str | int | None) -> str:
    """Return a deterministic synthetic key for V2 closing packages.

    Older logic relied on real WorkPackage ORM objects. In the V2 flow we only
    need a stable identifier to tag legacy TaskORM records so that the
    front-end can highlight which closing package they belong to. The exact
    format is not critical as long as it remains stable for the same pair of
    identifiers.
    """

    if package_id is None:
        return "package:unknown"

    if performer_id is None:
        return f"package:{package_id}"

    return f"package:{package_id}:performer:{performer_id}"
from decimal import Decimal, ROUND_HALF_UP
from pathlib import Path
from html import escape
from html.parser import HTMLParser


from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell
from sqlalchemy import select
from sqlalchemy.orm import Session, selectinload, object_session
from sqlalchemy.orm.attributes import flag_modified

from .. import orm_models
from .directory import ensure_individual_account
from ..config import BASE_DIR
from ..schemas import (
    DocumentApprovalAction,
    DocumentApprovalNote,
    DocumentNoteRequest,
    DocumentCreateRequest,
    DocumentCreateResponse,
    DocumentFile,
    DocumentRecord,
    TaskCostSnapshot,
    WorkPackage,
    WorkPackageMetadata,
    DocumentAssignee,
    UserPublic,
)
from ..storage import get_template

DOCUMENTS_DIR = BASE_DIR / "data" / "documents"
DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)

DOCUMENT_TYPE_MAP = {
    "Акт": "act",
    "Счет": "invoice",
    "Пакет": "package",
    "Приказ": "order",
}

DOC_V2_TYPE_LABELS = {
    "AVR": "Акт",
    "APP": "Акт",
    "IPR": "Акт",
    "INVOICE": "Счет",
    "SERVICE_ASSIGN": "Служебное задание",
    "ORDER": "Приказ",
}

DOC_V2_FILE_TYPES = {
    "AVR": "act",
    "APP": "act",
    "IPR": "act",
    "INVOICE": "invoice",
    "SERVICE_ASSIGN": "internal",
    "ORDER": "internal",
}

DOC_V2_PRIORITY = {
    "AVR": 0,
    "APP": 1,
    "IPR": 2,
    "INVOICE": 3,
    "SERVICE_ASSIGN": 4,
    "ORDER": 5,
}

DOC_V2_AUDIENCE = {
    "AVR": ["act"],
    "APP": ["act"],
    "IPR": ["act"],
    "INVOICE": ["invoice"],
    "SERVICE_ASSIGN": ["internal"],
    "ORDER": ["internal"],
}


PLACEHOLDER_PATTERN = re.compile(r"\$\{([^}]+)\}")


TEMPLATE_VARIABLE_EXPORT_KEYS: set[str] = {
    "companyName",
    "companyInn",
    "companyKpp",
    "companyBasis",
    "seoFullName",
    "seoShortName",
    "seoPosition",
    "seoAuthority",
    "responsiblePerson",
    "contractorCompanyName",
    "contractorSeoFullName",
    "contractorseoShortName",
    "employeeName",
    "employeeInn",
    "employeeContractNumber",
    "employeeContractDate",
    "assignmentGoal",
    "assignmentPurpose",
    "assignmentRequirements",
    "assignmentAppendix",
    "assignmentBasis",
    "appendixName",
    "deadlineDate",
    "orderNumber",
    "orderDate",
    "orderNumber1",
    "orderDate1",
    "orderNumber2",
    "orderDate2",
    "softwareName",
    "softwareCustomer",
    "softwareFunctionality",
    "projectSystemName",
    "repositorySystem",
    "devServer",
    "gitlabServer",
    "gptBody",
    "bodygpt",
    "actNumber",
    "startPeriodDate",
    "endPeriodDate",
    "totalAmountNumeric",
    "totalAmountWords",
    "vatAmountNumeric",
    "vatAmountWords",
}


CONTRACT_EXTRA_MAPPING: dict[str, str] = {
    "responsible_person": "responsiblePerson",
    "seo_full_name": "seoFullName",
    "seo_short_name": "seoShortName",
    "seo_position": "seoPosition",
    "seo_authority": "seoAuthority",
    "appendix_name": "appendixName",
    "assignment_goal": "assignmentGoal",
    "assignment_purpose": "assignmentPurpose",
    "assignment_requirements": "assignmentRequirements",
    "assignment_basis": "assignmentBasis",
    "assignment_appendix": "assignmentAppendix",
    "deadline_date": "deadlineDate",
    "order_number": "orderNumber",
    "order_date": "orderDate",
    "order_number_1": "orderNumber1",
    "order_date_1": "orderDate1",
    "order_number_2": "orderNumber2",
    "order_date_2": "orderDate2",
    "software_name": "softwareName",
    "software_customer": "softwareCustomer",
    "software_functionality": "softwareFunctionality",
    "project_system_name": "projectSystemName",
    "repository_system": "repositorySystem",
    "dev_server": "devServer",
    "gitlab_server": "gitlabServer",
}


logger = logging.getLogger(__name__)


def _normalize_v2_status(value: object) -> str:
    if value is None:
        return "draft"
    text = str(value).strip().lower()
    if not text:
        return "draft"
    text = text.replace('-', '_').replace(' ', '_')
    mapping = {
        'draft': 'draft',
        'pending_performer': 'pending_performer',
        'pendingperformer': 'pending_performer',
        'pending_performer_approval': 'pending_performer',
        'pending_manager': 'pending_manager',
        'pendingmanager': 'pending_manager',
        'pending_manager_approval': 'pending_manager',
        'manager_approved': 'manager_approved',
        'managerapproved': 'manager_approved',
        'manager-approved': 'manager_approved',
        'managerapprovedpending': 'manager_approved',
        'performer_approved': 'pending_manager',
        'performerapproved': 'pending_manager',
        'performer-approved': 'pending_manager',
        'rejected_performer': 'rejected_performer',
        'performer_rejected': 'rejected_performer',
        'performer-rejected': 'rejected_performer',
        'rejected_manager': 'rejected_manager',
        'manager_rejected': 'rejected_manager',
        'manager-rejected': 'rejected_manager',
        'final': 'final',
        'finalized': 'final',
        'completed': 'final',
    }
    return mapping.get(text, 'draft')


def _format_currency(value: float, currency: str = "RUB") -> str:
    quantized = Decimal(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    formatted = f"{quantized:,.2f}".replace(",", " ")
    return f"{formatted} {currency}"


def _format_hours(value: float) -> str:
    quantized = Decimal(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{quantized}".replace(",", ".")


def _format_date(value: datetime | None) -> str:
    if not value:
        return ""
    return value.strftime("%d.%m.%Y")


def _render_template_content(content: str, context: dict[str, str]) -> str:
    def replace(match: re.Match[str]) -> str:
        key = match.group(1)
        return str(context.get(key, ""))

    return PLACEHOLDER_PATTERN.sub(replace, content)


def _stringify_value(value: object) -> str:
    if isinstance(value, bool):
        return "Да" if value else "Нет"
    if isinstance(value, (int, float, Decimal)):
        return str(value)
    if isinstance(value, datetime):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, date):
        return value.strftime("%d.%m.%Y")
    if isinstance(value, str):
        text = value.strip()
        if re.fullmatch(r"\d{4}-\d{2}-\d{2}", text):
            try:
                parsed = datetime.strptime(text, "%Y-%m-%d")
                return parsed.strftime("%d.%m.%Y")
            except ValueError:
                return text
        return text
    return str(value)


def _short_name(full_name: str | None) -> str:
    if not full_name:
        return ""
    parts = re.split(r"[\s,]+", full_name.strip())
    parts = [part for part in parts if part]
    if not parts:
        return ""
    if len(parts) == 1:
        return parts[0]
    initials = [f"{segment[0].upper()}." for segment in parts[1:] if segment]
    initials_text = " ".join(initials[:2])
    return f"{parts[0]} {initials_text}".strip()


def _split_position_authority(basis: str | None) -> tuple[str | None, str | None]:
    if not basis:
        return (None, None)
    text = basis.strip()
    if not text:
        return (None, None)
    parts = [item.strip() for item in text.split(",", 1) if item]
    if len(parts) == 2:
        return (parts[0] or None, parts[1] or None)
    match = re.search(r"(действующ[\w\s]+на\s+основании.*)", text, re.IGNORECASE)
    if match:
        authority = match.group(1).strip()
        position = text[: match.start()].strip(" ,;")
        return (position or None, authority or None)
    return (None, text)


def _parse_period_bounds(period: str | None) -> tuple[str | None, str | None]:
    if not period:
        return (None, None)
    match = re.search(r"(20\d{2})[\.\-/](\d{1,2})", period)
    if not match:
        return (None, None)
    year = int(match.group(1))
    month = int(match.group(2))
    if not (1 <= month <= 12):
        return (None, None)
    first_day = date(year, month, 1)
    last_day = date(year, month, monthrange(year, month)[1])
    return (first_day.strftime("%d.%m.%Y"), last_day.strftime("%d.%m.%Y"))


def _sanitize_doc_number(value: str | None) -> str:
    if not value:
        return ""
    cleaned = re.sub(r"[^0-9A-Za-z\-/]", "", value.upper())
    return cleaned.strip("-/")


def _merge_template_variables(context: dict[str, str], *sources: dict[str, object], overwrite: bool = False) -> None:
    for source in sources:
        if not isinstance(source, dict):
            continue
        for key, raw_value in source.items():
            if raw_value is None:
                continue
            text = _stringify_value(raw_value)
            if isinstance(text, str) and not text.strip():
                continue
            if overwrite or not context.get(key):
                context[key] = text


def _collect_client_variables(client: orm_models.LegalEntityORM | None) -> dict[str, object]:
    if not client:
        return {}
    raw_basis = (client.basis or "").strip()
    position, authority = _split_position_authority(raw_basis)
    if authority and authority.lower() == "другое" and raw_basis:
        authority = raw_basis
    if not authority and raw_basis:
        authority = raw_basis
    if authority:
        authority = authority.strip()
    result: dict[str, object] = {
        "companyName": client.name or "",
        "companyInn": client.inn or "",
        "companyKpp": client.kpp or "",
        "seoFullName": client.signatory or "",
        "seoShortName": _short_name(client.signatory),
        "seoPosition": position or "",
        "seoAuthority": authority or raw_basis,
        "companyBasis": raw_basis,
        "powerOfAttorneyNumber": client.power_of_attorney_number or None,
        "powerOfAttorneyDate": client.power_of_attorney_date or None,
        "responsiblePerson": client.signatory or "",
    }
    # legacy aliases
    result["clientInn"] = result.get("companyInn", "")
    result["clientKpp"] = result.get("companyKpp", "")
    if not result.get("seoPosition"):
        result["seoPosition"] = "Генерального директора"
    return {key: value for key, value in result.items() if value not in (None, "")}


def _collect_contractor_variables(contractor: orm_models.IndividualORM | None, performer_type: str | None) -> dict[str, object]:
    if not contractor:
        return {}
    performer_name = contractor.name or ""
    short_name = _short_name(performer_name)
    result: dict[str, object] = {
        "employeeName": performer_name,
        "employeeInn": contractor.inn or "",
        "employeeAddress": contractor.address or "",
        "employeeEmail": contractor.email or "",
        "contractorCompanyName": performer_name,
        "contractorSeoFullName": performer_name,
        "contractorseoShortName": short_name,
    }
    if performer_type == "employee" and performer_name:
        result.setdefault("responsiblePerson", performer_name)
    return {key: value for key, value in result.items() if value not in (None, "")}


def _collect_contract_variables(contract: orm_models.ContractORM | None) -> dict[str, object]:
    if not contract:
        return {}
    result: dict[str, object] = {
        "contractNumber": contract.number or "",
        "employeeContractNumber": contract.number or "",
    }
    contract_date = getattr(contract, "contract_date", None) or getattr(contract, "created_at", None)
    if isinstance(contract_date, (datetime, date)):
        result["employeeContractDate"] = contract_date

    settings = getattr(contract, "settings", None)
    if settings and getattr(settings, "extra", None):
        extra = settings.extra or {}
        if isinstance(extra, dict):
            for src_key, dst_key in CONTRACT_EXTRA_MAPPING.items():
                if src_key in extra and extra[src_key] not in (None, ""):
                    result[dst_key] = extra[src_key]
            nested_candidates = [
                extra.get("template_variables"),
                extra.get("templateVariables"),
                extra.get("document_variables"),
            ]
            for candidate in nested_candidates:
                if isinstance(candidate, dict):
                    for key, value in candidate.items():
                        if value not in (None, ""):
                            result[str(key)] = value
    return {key: value for key, value in result.items() if value not in (None, "")}


def _collect_contract_meta_variables(meta: dict | None) -> dict[str, object]:
    if not isinstance(meta, dict):
        return {}
    result: dict[str, object] = {}
    for src_key, dst_key in CONTRACT_EXTRA_MAPPING.items():
        value = meta.get(src_key)
        if value in (None, ""):
            continue
        result[dst_key] = value
    if 'orderNumber' in result:
        result.setdefault('orderNumber1', result['orderNumber'])
    if 'orderDate' in result:
        result.setdefault('orderDate1', result['orderDate'])
    if not result.get("seoPosition"):
        result["seoPosition"] = "Генерального директора"
    nested_candidates = [
        meta.get("template_variables"),
        meta.get("templateVariables"),
        meta.get("document_variables"),
    ]
    for candidate in nested_candidates:
        if isinstance(candidate, dict):
            for key, value in candidate.items():
                if value in (None, ""):
                    continue
                result[str(key)] = value
    return result


def _collect_metadata_variables(metadata: object) -> dict[str, object]:
    if not isinstance(metadata, dict):
        return {}
    variables: dict[str, object] = {}
    for key in ("template_variables", "templateVariables", "variables"):
        candidate = metadata.get(key)
        if isinstance(candidate, dict):
            for inner_key, value in candidate.items():
                if value not in (None, ""):
                    variables[str(inner_key)] = value
    return variables


def _derive_act_number(
    *,
    context: dict[str, str],
    work_package: orm_models.WorkPackageORM,
    contract: orm_models.ContractORM | None,
    payload: DocumentCreateRequest,
) -> str:
    if context.get("actNumber"):
        return context["actNumber"]
    parts: list[str] = []
    if contract and contract.number:
        parts.append(_sanitize_doc_number(contract.number))
    period = payload.period or work_package.period or ""
    start_label, end_label = _parse_period_bounds(period)
    if start_label and end_label:
        period_token = f"{start_label[:2]}{start_label[3:5]}-{end_label[:2]}{end_label[3:5]}"
        parts.append(period_token)
    elif period:
        parts.append(_sanitize_doc_number(period))
    parts.append(work_package.id[-6:].upper())
    generated = "-".join(filter(None, parts))
    return _sanitize_doc_number(generated)


def _task_attribute(source, *names):
    for name in names:
        if isinstance(source, dict):
            if name in source and source.get(name) is not None:
                return source.get(name)
        else:
            value = getattr(source, name, None)
            if value is not None:
                return value
    return None


def _prepare_task_bullets(items: Iterable) -> list[dict[str, object]]:
    bullets: list[dict[str, object]] = []
    for task in items:
        hours_raw = _task_attribute(task, 'hours')
        try:
            hours = float(hours_raw) if hours_raw is not None else 0.0
        except (TypeError, ValueError):
            hours = 0.0
        force_included = bool(_task_attribute(task, "force_included", "forceIncluded") or False)
        description = (_task_attribute(task, "description", "body", "details", "content") or "")
        bullets.append(
            {
                "key": _issue_key(str(_task_attribute(task, "id", "key") or "")),
                "summary": (_task_attribute(task, "summary", "title") or ""),
                "description": description or "",
                "hours": round(hours, 2),
                "forceIncluded": force_included,
            }
        )
    return bullets


def _strip_html_markup(value: str) -> str:
    text = value or ""
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"(?i)</(p|div|li|h[1-6]|br|tr|table)>", "\n", text)
    text = text.replace("&lt;", "<").replace("&gt;", ">").replace("&amp;", "&")
    text = re.sub(r"<[^>]+>", "", text)
    text = re.sub(r"^\s*#{1,6}\s*", "", text, flags=re.M)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    text = re.sub(r"^\s*([-*]|\d+\.)\s+", "", text, flags=re.M)
    text = re.sub(r"[ \t]+$", "", text, flags=re.M)
    return text.strip()


def _strip_meta_lines(value: str) -> str:
    if not value:
        return value
    text = re.sub(r"^\s*#{0,6}\s*Акт\s+выполненных\s+работ[^\n]*\n?", "", value, flags=re.I | re.M)
    text = re.sub(r"^\s*#{0,6}\s*Statement\s+of\s+Work[^\n]*\n?", "", text, flags=re.I | re.M)
    text = re.sub(r"^\s*(Проект:|Ключ проекта:|Период выполнения работ:|Общее количество часов:|Выполненные задачи|Примечания)\s*:?\s*.*\n?", "", text, flags=re.I | re.M)
    text = re.sub(r"^\s*(Project:|Project key:|Period:|Total hours:|Tasks|Notes)\s*:?\s*.*\n?", "", text, flags=re.I | re.M)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]{2,}", " ", text)
    return text.strip()


def _shorten(value: str, limit: int = 1500) -> str:
    value = (value or "").strip()
    return (value[:limit] + "…") if len(value) > limit else value


def _normalize_bullet_sentence(value: str) -> str:
    if not value:
        return ""
    text = _strip_meta_lines(_strip_html_markup(value))
    text = re.sub(r"\[[^\]]*\]\s*", "", text)
    text = re.sub(r"ACS-\d+\s*[—\-:]\s*", "", text, flags=re.I)
    text = re.sub(r"^[\s\d\.\)\-–—•*]+", "", text)
    text = re.sub(r"[ \t]+", " ", text).strip()
    if text and not text.endswith("."):
        text = f"{text}."
    return text


def _build_plain_act_html(bullets: list[dict[str, object]], *, lang: str) -> str:
    filtered = [
        item
        for item in bullets
        if (item.get("description") and str(item["description"]).strip()) or item.get("forceIncluded")
    ]
    if not filtered:
        filtered = bullets

    sentences: list[str] = []
    seen = set()
    for bullet in filtered:
        raw = bullet.get("description") or bullet.get("summary") or ""
        normalized = _normalize_bullet_sentence(str(raw))
        if not normalized:
            continue
        shortened = _shorten(normalized)
        if shortened in seen:
            continue
        seen.add(shortened)
        sentences.append(shortened)

    if not sentences:
        return ""

    limited = sentences[:6]
    if len(sentences) > 6:
        limited.append(
            "Дополнительно выполнены задачи, отражённые в таблице работ."
            if lang != "en"
            else "Additional tasks are listed in the work log table."
        )

    intro = (
        "During the reporting period the following work was completed:"
        if lang == "en"
        else "В отчетный период были выполнены следующие работы:"
    )
    body = " ".join(limited)

    combined = f"{intro} {body}".strip()
    return "\n".join(
        [
            '<div class="doc-template doc-template--act">',
            f"<p>{escape(combined)}</p>",
            "</div>",
        ]
    )


def _build_tasks_table(task_snapshots: Iterable[dict]) -> str:
    rows: list[str] = []
    for snapshot in task_snapshots:
        title = f"{snapshot.get('title', '')} ({snapshot.get('key', '')})"
        hours = _format_hours(float(snapshot.get('hours', 0)))
        rate = _format_currency(float(snapshot.get('hourlyRate', 0)))
        amount = _format_currency(float(snapshot.get('amount', 0)))
        rows.append(
            "<tr>"
            f"<td>{escape(title)}</td>"
            f"<td>{escape(hours)}</td>"
            f"<td>{escape(rate)}</td>"
            f"<td>{escape(amount)}</td>"
            "</tr>"
        )

    if not rows:
        rows.append('<tr><td colspan="4">—</td></tr>')

    return ''.join(rows)


def _create_docx_from_text(content: str, destination: Path) -> None:
    document = Document()
    _clear_document(document)
    _initialize_document_defaults(document)
    renderer = DocxHtmlRenderer(document)
    renderer.render(content)
    destination.parent.mkdir(parents=True, exist_ok=True)
    document.save(destination)


def _clean_string(value: str | None) -> str | None:
    if value is None:
        return None
    cleaned = value.strip()
    return cleaned or None


def _unique_strings(values: Iterable[str | None]) -> list[str]:
    unique: dict[str, None] = {}
    for value in values:
        if value is None:
            continue
        cleaned = value.strip()
        if not cleaned:
            continue
        unique.setdefault(cleaned, None)
    return list(unique.keys())


def _compose_metadata(
    *,
    existing: dict | None,
    payload: DocumentCreateRequest,
    include_timesheet: bool,
    document_type: str,
    period: str,
    project_key: str | None,
) -> dict:
    metadata = dict(existing or {})

    source_prepared = payload.audience or metadata.get("preparedFor") or ["act", "invoice", "tax-report"]
    prepared_for = [
        item
        for item in source_prepared
        if isinstance(item, str) and item.strip()
    ]
    metadata["preparedFor"] = prepared_for

    doc_tag = DOCUMENT_TYPE_MAP.get(document_type, document_type)
    tags = _unique_strings(
        [
            *(metadata.get("tags") or []),
            *(payload.tags or []),
            *prepared_for,
            period,
            project_key,
            doc_tag,
            document_type,
            "timesheet" if include_timesheet else None,
        ]
    )
    metadata["tags"] = tags

    if payload.taxCategory is not None:
        metadata["taxCategory"] = _clean_string(payload.taxCategory)
    elif not metadata.get("taxCategory") and "tax-report" in prepared_for:
        metadata["taxCategory"] = "Налоговая отчётность"

    if payload.benefitCategory is not None:
        metadata["benefitCategory"] = _clean_string(payload.benefitCategory)
    elif not metadata.get("benefitCategory") and "benefit-report" in prepared_for:
        metadata["benefitCategory"] = "Льготы / субсидии"

    metadata["currency"] = metadata.get("currency") or "RUB"

    return metadata


def _clear_document(document: Document) -> None:
    while len(document.paragraphs) > 1:
        _remove_paragraph(document.paragraphs[-1])
    if document.paragraphs:
        _clear_paragraph(document.paragraphs[0])
    else:
        document.add_paragraph()


def _initialize_document_defaults(document: Document) -> None:
    try:
        normal_style = document.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        para = normal_style.paragraph_format
        para.line_spacing = 1.2
        para.space_after = Pt(6)
        para.first_line_indent = Pt(_css_length_to_pt('1.25cm') or 35.4)
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    except KeyError:
        pass


class DocxHtmlRenderer(HTMLParser):
    def __init__(self, document: Document, container: _Cell | None = None) -> None:
        super().__init__(convert_charrefs=True)
        self.document = document
        self.container = container
        self.current_paragraph = None
        self.run_attrs_stack: list[dict[str, bool]] = [{
            'bold': False,
            'italic': False,
            'underline': False,
            'color': None,
            'size': None,
            'font': None,
        }]
        self.list_stack: list[str] = []
        self.table_stack: list[dict[str, object]] = []
        self.capture_stack: list[dict[str, object]] = []
        self.cell_initialized = False
        self.initial_paragraph = None
        self.initial_paragraph_used = False
        self._paragraph_style_depth = 0
        self._current_paragraph_classes: set[str] = set()
        self.section_stack: list[dict[str, object]] = []
        self.list_item_stack: list[dict[str, object]] = []
        self._paragraph_reuse_stack: list[bool] = []

        if container is None:
            if document.paragraphs:
                self.initial_paragraph = document.paragraphs[0]
                _clear_paragraph(self.initial_paragraph)
        else:
            if container.paragraphs:
                _clear_paragraph(container.paragraphs[0])

    # Public interface --------------------------------------------------

    def render(self, html: str) -> None:
        if not html:
            return
        self.feed(html)
        self.close()
        self.finalize()

    def finalize(self) -> None:
        self.current_paragraph = None

    # HTMLParser overrides ----------------------------------------------

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:  # type: ignore[override]
        tag_lower = tag.lower()

        if self.capture_stack:
            self.capture_stack[-1]['buffer'].append(_format_start_tag(tag, attrs))
            self.capture_stack[-1]['depth'] += 1
            return

        if self._capture_start(tag_lower, attrs):
            return

        attr_map = _attrs_to_dict(attrs)
        style_map = _extract_style(attrs)
        element_classes = _extract_classes(attr_map)
        if tag_lower == 'div':
            self.section_stack.append({'classes': element_classes.copy(), 'paragraph_count': 0, 'table_count': 0})

        inherited_classes: set[str] = set()
        for context in self.section_stack:
            inherited_classes.update(context.get('classes', set()))
        classes = element_classes | inherited_classes

        _apply_class_to_style(style_map, classes)

        if tag_lower == 'div' and element_classes.intersection({'doc-template', 'doc-template-preview'}):
            return

        if tag_lower in {'p', 'div'}:
            reuse = False
            if tag_lower == 'p' and self.list_item_stack:
                li_ctx = self.list_item_stack[-1]
                paragraph = li_ctx.get('paragraph')
                if paragraph is not None:
                    if self.current_paragraph is None:
                        self.current_paragraph = paragraph  # reuse list item paragraph
                    reuse = True
            self._start_paragraph(attr_map, style_map, classes, reuse_current=reuse)
        elif tag_lower in {'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            para = self._start_paragraph(attr_map, style_map, classes)
            try:
                para.style = f"Heading {int(tag_lower[1])}"
            except (ValueError, KeyError):
                pass
        elif tag_lower in {'strong', 'b'}:
            self._push_run_attrs(bold=True)
        elif tag_lower in {'em', 'i'}:
            self._push_run_attrs(italic=True)
        elif tag_lower in {'u', 'ins'}:
            self._push_run_attrs(underline=True)
        elif tag_lower == 'span':
            if 'doc-flexline__tab' in element_classes:
                paragraph = self._ensure_paragraph()
                paragraph.add_run().add_tab()
            self._push_run_attrs(style_map=style_map)
        elif tag_lower == 'br':
            paragraph = self._ensure_paragraph()
            paragraph.add_run().add_break()
        elif tag_lower == 'ul':
            self.list_stack.append('bullet')
        elif tag_lower == 'ol':
            self.list_stack.append('number')
        elif tag_lower == 'li':
            list_type = self.list_stack[-1] if self.list_stack else 'bullet'
            paragraph = self._start_paragraph(attr_map, style_map, classes, list_type=list_type)
            self.list_item_stack.append({'paragraph': paragraph})
        elif tag_lower == 'table':
            self._break_paragraph()
            self.table_stack.append({'rows': [], 'current_row': None, 'attrs': attr_map, 'style': style_map, 'classes': classes})
        elif tag_lower == 'tr':
            if self.table_stack:
                self.table_stack[-1]['current_row'] = []

    def handle_endtag(self, tag: str) -> None:  # type: ignore[override]
        tag_lower = tag.lower()

        if self._capture_end(tag_lower):
            return

        if tag_lower in {'strong', 'b', 'em', 'i', 'u', 'ins', 'span'}:
            if len(self.run_attrs_stack) > 1:
                self.run_attrs_stack.pop()
        elif tag_lower in {'p', 'div', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6'}:
            self._break_paragraph()
            if tag_lower == 'li' and self.list_item_stack:
                self.list_item_stack.pop()
        elif tag_lower in {'ul', 'ol'}:
            if self.list_stack:
                self.list_stack.pop()
        elif tag_lower == 'tr':
            if self.table_stack:
                ctx = self.table_stack[-1]
                row = ctx.get('current_row')
                if row is not None:
                    ctx['rows'].append(row)
                    ctx['current_row'] = None
        elif tag_lower == 'table':
            if self.table_stack:
                ctx = self.table_stack.pop()
                self._flush_table(ctx)
        elif tag_lower == 'div':
            if self.section_stack:
                self.section_stack.pop()

    def handle_data(self, data: str) -> None:  # type: ignore[override]
        if self.capture_stack:
            self.capture_stack[-1]['buffer'].append(data)
            return
        self._append_text(data)

    # Capture utilities -------------------------------------------------

    def _capture_start(self, tag: str, attrs: list[tuple[str, str | None]]) -> bool:
        if not self.table_stack or tag not in {'td', 'th'}:
            return False
        ctx = self.table_stack[-1]
        if ctx.get('current_row') is None:
            ctx['current_row'] = []
        attr_map = _attrs_to_dict(attrs)
        class_set = _extract_classes(attr_map)
        cell_ref = {'tag': tag, 'html': '', 'attrs': attr_map, 'classes': class_set}
        ctx['current_row'].append(cell_ref)
        self.capture_stack.append({
            'tag': tag,
            'buffer': [_format_start_tag(tag, attrs)],
            'depth': 1,
            'context': ctx,
            'cell': cell_ref,
        })
        return True

    def _capture_end(self, tag: str) -> bool:
        if not self.capture_stack:
            return False
        capture = self.capture_stack[-1]
        capture['buffer'].append(f"</{tag}>")
        capture['depth'] -= 1
        if capture['depth'] <= 0:
            cell_ref = capture['cell']
            cell_ref['html'] = ''.join(capture['buffer'])
            self.capture_stack.pop()
        return True

    # Paragraph helpers -------------------------------------------------

    def _start_paragraph(
        self,
        attr_map: dict[str, str | None] | None,
        style_map: dict[str, str],
        classes: set[str],
        *,
        list_type: str | None = None,
        reuse_current: bool = False,
    ):
        style_map = dict(style_map)
        for context in reversed(self.section_stack):
            context['paragraph_count'] = context.get('paragraph_count', 0) + 1
            _apply_section_paragraph_rules(style_map, context.get('classes', set()), context['paragraph_count'])

        paragraph = None
        if reuse_current and self.current_paragraph is not None:
            paragraph = self.current_paragraph
        else:
            paragraph = self._new_paragraph(list_type)
        if reuse_current:
            self._current_paragraph_classes.update(classes)
        else:
            self._current_paragraph_classes = set(classes)
        alignment = None
        if attr_map and attr_map.get('align'):
            alignment = (attr_map.get('align') or '').lower()
        if style_map.get('text-align'):
            alignment = style_map['text-align']
        if alignment:
            if alignment == 'center':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif alignment == 'right':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif alignment == 'justify':
                paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _apply_paragraph_styling(paragraph, style_map, classes)
        self.current_paragraph = paragraph
        self._paragraph_style_depth += 1
        self._push_run_attrs(style_map=style_map)
        self._paragraph_reuse_stack.append(reuse_current)
        return paragraph

    def _new_paragraph(self, list_type: str | None = None):
        style_name = None
        if list_type == 'bullet':
            style_name = 'List Bullet'
        elif list_type == 'number':
            style_name = 'List Number'

        if isinstance(self.container, _Cell):
            if not self.cell_initialized:
                paragraph = self.container.paragraphs[0]
                _clear_paragraph(paragraph)
                self.cell_initialized = True
            else:
                paragraph = self.container.add_paragraph()
        else:
            if self.initial_paragraph is not None and not self.initial_paragraph_used:
                paragraph = self.initial_paragraph
                self.initial_paragraph_used = True
            else:
                paragraph = self.document.add_paragraph()

        if style_name:
            try:
                paragraph.style = style_name
            except KeyError:
                pass
        return paragraph

    def _ensure_paragraph(self):
        if self.current_paragraph is None:
            self.current_paragraph = self._new_paragraph()
        return self.current_paragraph

    def _break_paragraph(self) -> None:
        reuse = self._paragraph_reuse_stack.pop() if self._paragraph_reuse_stack else False
        if self._paragraph_style_depth > 0:
            if len(self.run_attrs_stack) > 1:
                self.run_attrs_stack.pop()
            self._paragraph_style_depth -= 1
        if reuse:
            return
        self.current_paragraph = None
        self._current_paragraph_classes = set()

    # Formatting helpers ------------------------------------------------

    def _push_run_attrs(
        self,
        *,
        bold: bool | None = None,
        italic: bool | None = None,
        underline: bool | None = None,
        style_map: dict[str, str] | None = None,
    ) -> None:
        attrs = self.run_attrs_stack[-1].copy()
        if bold is not None:
            attrs['bold'] = bold
        if italic is not None:
            attrs['italic'] = italic
        if underline is not None:
            attrs['underline'] = underline
        if style_map:
            font_weight = style_map.get('font-weight')
            if font_weight in {'bold', '700'}:
                attrs['bold'] = True
            font_style = style_map.get('font-style')
            if font_style == 'italic':
                attrs['italic'] = True
            text_decoration = style_map.get('text-decoration')
            if text_decoration and 'underline' in text_decoration:
                attrs['underline'] = True
            color_value = style_map.get('color')
            if color_value:
                attrs['color'] = color_value
            font_size = style_map.get('font-size')
            if font_size:
                attrs['size'] = font_size
            font_family = style_map.get('font-family')
            if font_family:
                attrs['font'] = font_family
        self.run_attrs_stack.append(attrs)

    def _append_text(self, text: str) -> None:
        if not text:
            return
        paragraph = self._ensure_paragraph()
        text = text.replace('\r', '')
        parts = text.split('\n')
        emitted = bool(paragraph.text)
        for index, part in enumerate(parts):
            if index > 0 and emitted:
                paragraph.add_run().add_break()
            elif index > 0 and not emitted and not part:
                continue
            if part:
                run = paragraph.add_run(_preserve_spaces(part))
                self._apply_run_attrs(run)
                emitted = True

    def _apply_run_attrs(self, run) -> None:
        attrs = self.run_attrs_stack[-1]
        run.bold = True if attrs.get('bold') else None
        run.italic = True if attrs.get('italic') else None
        run.underline = True if attrs.get('underline') else None
        color_value = attrs.get('color')
        if color_value:
            rgb = _parse_color(color_value)
            if rgb is not None:
                run.font.color.rgb = rgb
        size_value = attrs.get('size')
        if size_value:
            size_pt = _css_length_to_pt(size_value)
            if size_pt is not None:
                run.font.size = Pt(size_pt)
        font_value = attrs.get('font')
        if font_value:
            family = _parse_font_family(font_value)
            if family:
                run.font.name = family

    def _flush_table(self, ctx: dict[str, object]) -> None:
        rows = _normalize_rows(ctx.get('rows') or [])
        if not rows:
            return
        max_cols = max(len(row) for row in rows)
        table = self.document.add_table(rows=len(rows), cols=max_cols)
        table_classes: set[str] = set()
        attrs = ctx.get('attrs')
        if isinstance(attrs, dict):
            table_classes.update(_extract_classes(attrs))
        stack_classes = ctx.get('classes')
        if isinstance(stack_classes, set):
            table_classes.update(stack_classes)
        for section in self.section_stack:
            table_classes.update(section.get('classes', set()))
        _apply_table_class_style(table, table_classes)
        for row_idx, row in enumerate(rows):
            for col_idx in range(max_cols):
                cell = table.cell(row_idx, col_idx)
                _clear_paragraph(cell.paragraphs[0])
                if col_idx < len(row):
                    cell_ref = row[col_idx]
                    cell_html = cell_ref.get('html') if isinstance(cell_ref, dict) else ''
                    renderer = DocxHtmlRenderer(self.document, container=cell)
                    renderer.render(cell_html)
                    if isinstance(cell_ref, dict) and cell_ref.get('tag') == 'th':
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True
        self._break_paragraph()
        for section in reversed(self.section_stack):
            section['table_count'] = section.get('table_count', 0) + 1
            _apply_section_table_rules(table, section.get('classes', set()), section['table_count'])


def _extract_style(attrs: list[tuple[str, str | None]]) -> dict[str, str]:
    style_text = ''
    for name, value in attrs:
        if name == 'style' and value:
            style_text = value
            break
    styles: dict[str, str] = {}
    for chunk in style_text.split(';'):
        if ':' in chunk:
            key, val = chunk.split(':', 1)
            styles[key.strip().lower()] = val.strip().lower()
    return styles


def _attrs_to_dict(attrs: list[tuple[str, str | None]]) -> dict[str, str | None]:
    return {name: value for name, value in attrs}


def _clear_paragraph(paragraph) -> None:
    for run in paragraph.runs:
        run.text = ''
    paragraph.text = ''


def _format_start_tag(tag: str, attrs: list[tuple[str, str | None]]) -> str:
    parts = [tag]
    for name, value in attrs:
        if value is None:
            continue
        escaped = value.replace('"', '&quot;') if isinstance(value, str) else ''
        parts.append(f'{name}="{escaped}"')
    return '<' + ' '.join(parts) + '>'


def _remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _extract_classes(attr_map: dict[str, str | None] | None) -> set[str]:
    if not attr_map:
        return set()
    raw = attr_map.get('class')
    if not raw:
        return set()
    classes = {cls.strip() for cls in raw.replace('\t', ' ').split(' ') if cls.strip()}
    return classes


def _apply_class_to_style(style_map: dict[str, str], classes: set[str]) -> None:
    if 'doc-center' in classes or 'text-center' in classes:
        style_map.setdefault('text-align', 'center')
    if 'doc-right' in classes or 'text-right' in classes:
        style_map.setdefault('text-align', 'right')
    if 'doc-justify' in classes or 'text-justify' in classes:
        style_map.setdefault('text-align', 'justify')
    if 'doc-no-indent' in classes or 'no-indent' in classes:
        style_map.setdefault('text-indent', '0')
        style_map.setdefault('margin-left', '0')
    if 'doc-flexline' in classes:
        style_map.setdefault('margin-left', '0')
        style_map.setdefault('text-indent', '0')


def _apply_table_class_style(table, classes: set[str]) -> None:
    if {'doc-table-transparent', 'doc-table-plain'}.intersection(classes):
        _remove_table_borders(table)
    elif 'doc-table-signature' in classes:
        _set_table_borders(table, 1.5)
    elif 'doc-table-bordered' in classes or 'doc-table-striped' in classes:
        _set_table_borders(table, 1.0)


def _remove_table_borders(table) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top=None, left=None, bottom=None, right=None)


def _set_table_borders(table, width_pt: float) -> None:
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top=width_pt, left=width_pt, bottom=width_pt, right=width_pt)


def _normalize_rows(rows: list[list[dict]]) -> list[list[dict]]:
    normalized: list[list[dict]] = []
    for row in rows:
        expanded = _expand_row(row)
        normalized.extend(expanded)
    return normalized


def _expand_row(row: list[dict]) -> list[list[dict]]:
    for cell in row:
        if isinstance(cell, dict):
            html = cell.get('html', '')
            if html and '<tr' in html.lower():
                extracted = _extract_rows_from_fragment(html)
                if extracted:
                    return extracted
    return [row]


def _extract_rows_from_fragment(fragment: str) -> list[list[dict]]:
    rows: list[list[dict]] = []
    for row_match in re.finditer(r'<tr[^>]*>(.*?)</tr>', fragment, flags=re.IGNORECASE | re.DOTALL):
        row_html = row_match.group(0)
        cells: list[dict] = []
        for cell_match in re.finditer(r'<(td|th)([^>]*)>(.*?)</\1>', row_html, flags=re.IGNORECASE | re.DOTALL):
            tag = cell_match.group(1).lower()
            attrs_raw = cell_match.group(2) or ''
            inner_html = cell_match.group(3) or ''
            attr_map = _parse_html_attributes(attrs_raw)
            classes = _extract_classes(attr_map)
            cells.append({'tag': tag, 'html': inner_html, 'attrs': attr_map, 'classes': classes})
        if cells:
            rows.append(cells)
    return rows


def _parse_html_attributes(raw: str) -> dict[str, str]:
    attrs: dict[str, str] = {}
    for match in re.finditer(r'([a-zA-Z_:][-\w\.:]*)\s*=\s*("([^"]*)"|\'([^\']*)\'|([^\s"\'=<>`]+))', raw):
        name = match.group(1).lower()
        value = match.group(3) or match.group(4) or match.group(5) or ''
        attrs[name] = value
    return attrs


def _set_cell_border(cell, **kwargs) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        tc_pr.append(tc_borders)
    for edge, value in kwargs.items():
        element = tc_borders.find(qn(f'w:{edge}'))
        if element is None:
            element = OxmlElement(f'w:{edge}')
            tc_borders.append(element)
        if value is None:
            element.set(qn('w:val'), 'nil')
        else:
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), str(int(value * 8)))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), 'auto')


def _apply_paragraph_styling(paragraph, style_map: dict[str, str], classes: set[str]) -> None:
    fmt = paragraph.paragraph_format
    indent_value = style_map.get('margin-left') or style_map.get('padding-left')
    if indent_value:
        pt_value = _css_length_to_pt(indent_value)
        if pt_value is not None:
            fmt.left_indent = Pt(pt_value)

    text_indent = style_map.get('text-indent')
    if text_indent:
        pt_value = _css_length_to_pt(text_indent)
        if pt_value is not None:
            fmt.first_line_indent = Pt(pt_value)
    if not text_indent and ('doc-no-indent' in classes or 'no-indent' in classes or 'doc-flexline' in classes):
        fmt.first_line_indent = Pt(0)

    margin_top = style_map.get('margin-top')
    if margin_top:
        pt_value = _css_length_to_pt(margin_top)
        if pt_value is not None:
            fmt.space_before = Pt(pt_value)

    margin_bottom = style_map.get('margin-bottom')
    if margin_bottom:
        pt_value = _css_length_to_pt(margin_bottom)
        if pt_value is not None:
            fmt.space_after = Pt(pt_value)

    line_height = style_map.get('line-height')
    if line_height:
        line_height = line_height.strip()
        if line_height.endswith('%'):
            try:
                fmt.line_spacing = float(line_height[:-1]) / 100
            except ValueError:
                pass
        else:
            pt_value = _css_length_to_pt(line_height)
            if pt_value is not None:
                fmt.line_spacing = Pt(pt_value)
            else:
                try:
                    fmt.line_spacing = float(line_height)
                except ValueError:
                    pass

    if 'doc-flexline' in classes:
        fmt.tab_stops.clear_all()
        fmt.left_indent = Pt(0)
        fmt.first_line_indent = Pt(0)
        fmt.tab_stops.add_tab_stop(Pt(420), alignment=WD_TAB_ALIGNMENT.RIGHT)


def _apply_section_paragraph_rules(style_map: dict[str, str], classes: set[str], index: int) -> None:
    if 'doc-template--act' in classes:
        if index == 1:
            style_map.setdefault('text-align', 'center')
            style_map.setdefault('font-weight', 'bold')
            style_map.setdefault('font-size', '14pt')
            style_map.setdefault('text-indent', '0')
            style_map.setdefault('margin-left', '0')
            style_map.setdefault('margin-bottom', '4pt')
        elif index == 2:
            style_map.setdefault('text-align', 'center')
            style_map.setdefault('text-indent', '0')
            style_map.setdefault('margin-left', '0')
            style_map.setdefault('margin-bottom', '6pt')
        elif index == 3:
            style_map.setdefault('text-indent', '0')
            style_map.setdefault('margin-left', '0')
            style_map.setdefault('margin-bottom', '10pt')


def _apply_section_table_rules(table, classes: set[str], index: int) -> None:
    if 'doc-template--act' in classes:
        if index == 1:
            _set_table_spacing(table, before=16)


def _set_table_spacing(table, *, before: float | None = None, after: float | None = None) -> None:
    if before is not None and table.rows:
        for cell in table.rows[0].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(before)
    if after is not None and table.rows:
        for cell in table.rows[-1].cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(after)


def _css_length_to_pt(value: str | None) -> float | None:
    if not value:
        return None
    value = value.strip().lower()
    try:
        if value.endswith('pt'):
            return float(value[:-2])
        if value.endswith('px'):
            return float(value[:-2]) * 72 / 96
        if value.endswith('cm'):
            return float(value[:-2]) * 72 / 2.54
        if value.endswith('mm'):
            return float(value[:-2]) * 72 / 25.4
        if value.endswith('in'):
            return float(value[:-2]) * 72
        if value.endswith('pc'):
            return float(value[:-2]) * 12
        return float(value)
    except ValueError:
        return None


def _parse_font_family(value: str | None) -> str | None:
    if not value:
        return None
    families = [item.strip().strip('"').strip("'") for item in value.split(',')]
    for family in families:
        if family:
            return family
    return None


def _parse_color(value: str | None) -> RGBColor | None:
    if not value:
        return None
    value = value.strip().lower()
    if value.startswith('#'):
        hex_value = value[1:]
        if len(hex_value) == 3:
            hex_value = ''.join(ch * 2 for ch in hex_value)
        if len(hex_value) == 6:
            try:
                r = int(hex_value[0:2], 16)
                g = int(hex_value[2:4], 16)
                b = int(hex_value[4:6], 16)
                return RGBColor(r, g, b)
            except ValueError:
                return None
    if value.startswith('rgb') and value.endswith(')'):
        try:
            parts = value[value.find('(') + 1 : value.rfind(')')].split(',')
            r, g, b = [int(part.strip()) for part in parts[:3]]
            return RGBColor(r, g, b)
        except (ValueError, TypeError):
            return None
    NAMED = {
        'black': (0, 0, 0),
        'white': (255, 255, 255),
        'red': (255, 0, 0),
        'green': (0, 128, 0),
        'blue': (0, 0, 255),
        'gray': (128, 128, 128),
        'grey': (128, 128, 128),
    }
    if value in NAMED:
        r, g, b = NAMED[value]
        return RGBColor(r, g, b)
    return None


def _preserve_spaces(text: str) -> str:
    if not text:
        return text
    text = text.replace('\t', '    ')
    if text.startswith(' '):
        text = '\u00A0' + text[1:]
    if text.endswith(' '):
        text = text[:-1] + '\u00A0'
    while '  ' in text:
        text = text.replace('  ', ' \u00A0')
    return text


def _format_number_plain(value: float) -> str:
    quantized = Decimal(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    return f"{quantized:,.2f}".replace(',', ' ')


def _amount_to_words(value: float) -> str:
    quantized = Decimal(value).quantize(Decimal("0.01"), rounding=ROUND_HALF_UP)
    rubles = int(quantized)
    kopeks = int((quantized - rubles) * 100)

    rubles_words = _number_to_words_triplets(rubles, gender='masc', forms=('рубль', 'рубля', 'рублей'))
    kopeks_words = _number_to_words_triplets(kopeks, gender='fem', forms=('копейка', 'копейки', 'копеек'))

    if rubles_words:
        rubles_part = rubles_words
    else:
        rubles_part = 'ноль рублей'

    kopeks_part = kopeks_words or 'ноль копеек'

    result = f"{rubles_part} {kopeks_part}"
    return result.capitalize()


def _number_to_words_triplets(value: int, *, gender: str, forms: tuple[str, str, str]) -> str:
    if value == 0:
        return ''

    units = {
        'masc': ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять'],
        'fem': ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть', 'семь', 'восемь', 'девять'],
    }
    teens = ['десять', 'одиннадцать', 'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать', 'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    tens = ['', 'десять', 'двадцать', 'тридцать', 'сорок', 'пятьдесят', 'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот', 'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']
    scales = [
        (('рубль', 'рубля', 'рублей'), 'masc'),
        (('тысяча', 'тысячи', 'тысяч'), 'fem'),
        (('миллион', 'миллиона', 'миллионов'), 'masc'),
        (('миллиард', 'миллиарда', 'миллиардов'), 'masc'),
    ]

    words: list[str] = []
    remainder = value
    index = 0
    while remainder > 0:
        triplet = remainder % 1000
        remainder //= 1000
        if triplet == 0 and index > 0:
            index += 1
            continue
        current_gender = 'masc'
        current_forms = ('', '', '')
        if index == 0:
            current_gender = gender
            current_forms = forms
        else:
            if index < len(scales):
                current_forms, current_gender = scales[index]
            else:
                current_forms = (scales[-1][0][0], scales[-1][0][1], scales[-1][0][2])
                current_gender = 'masc'

        triplet_words = []
        h = triplet // 100
        t_u = triplet % 100
        t = t_u // 10
        u = t_u % 10

        if h:
            triplet_words.append(hundreds[h])
        if 10 <= t_u <= 19:
            triplet_words.append(teens[t_u - 10])
        else:
            if t:
                triplet_words.append(tens[t])
            if u:
                gender_units = units['fem'] if current_gender == 'fem' else units['masc']
                triplet_words.append(gender_units[u])

        if triplet_words:
            triplet_words.append(_choose_form(triplet, current_forms))
            words.insert(0, ' '.join(filter(None, triplet_words)))
        index += 1

    result = ' '.join(words).strip()
    return result


def _choose_form(value: int, forms: tuple[str, str, str]) -> str:
    value = abs(value) % 100
    if 10 < value < 20:
        return forms[2]
    value = value % 10
    if value == 1:
        return forms[0]
    if 2 <= value <= 4:
        return forms[1]
    return forms[2]


def _build_context(
    *,
    work_package: orm_models.WorkPackageORM,
    contract: orm_models.ContractORM | None,
    client: orm_models.LegalEntityORM | None,
    contractor: orm_models.IndividualORM | None,
    tasks: Iterable[orm_models.TaskORM],
    payload: DocumentCreateRequest,
    vat_amount: float,
) -> dict[str, str]:
    task_snapshots = work_package.task_snapshots or []
    total_amount = float(work_package.total_amount)
    total_hours = float(work_package.total_hours)
    vat_percent = float(payload.vatPercent or work_package.vat_percent or 0)

    if payload.vatIncluded or work_package.vat_included:
        without_vat = total_amount - vat_amount
    else:
        without_vat = total_amount

    task_dates = [task.updated_at for task in tasks if task.updated_at]
    start_date = min(task_dates).date() if task_dates else None
    end_date = max(task_dates).date() if task_dates else None

    context: dict[str, str] = {
        'date': _format_date(datetime.utcnow()),
        'projectName': work_package.project_name,
        'projectKey': work_package.project_key,
        'period': work_package.period,
        'totalHours': _format_hours(total_hours),
        'totalAmount': _format_currency(total_amount, work_package.currency),
        'totalAmountWithoutVat': _format_currency(without_vat, work_package.currency),
        'vatPercent': f"{vat_percent}",
        'vatAmount': _format_currency(vat_amount, work_package.currency),
        'hourlyRate': _format_currency(float(work_package.hourly_rate), work_package.currency),
        'tasksCount': str(len(task_snapshots)),
        'currency': work_package.currency,
        'table1': _build_tasks_table(task_snapshots),
        'tableTasks': _build_tasks_table(task_snapshots),
    }

    context['totalAmountNumeric'] = _format_number_plain(total_amount)
    context['totalAmountWords'] = _amount_to_words(total_amount)
    context['vatAmountNumeric'] = _format_number_plain(vat_amount)
    context['vatAmountWords'] = _amount_to_words(vat_amount)

    context['table2'] = context['tableTasks']

    metadata_map = work_package.metadata_json if isinstance(work_package.metadata_json, dict) else {}
    _merge_template_variables(context, _collect_metadata_variables(metadata_map))

    prepared_for = metadata_map.get('preparedFor') if isinstance(metadata_map, dict) else None
    if isinstance(prepared_for, list):
        context['preparedFor'] = ', '.join(prepared_for)

    if start_date:
        context['startPeriodDate'] = start_date.strftime('%d.%m.%Y')
    if end_date:
        context['endPeriodDate'] = end_date.strftime('%d.%m.%Y')

    if not context.get('startPeriodDate') or not context.get('endPeriodDate'):
        derived_start, derived_end = _parse_period_bounds(payload.period or work_package.period)
        if derived_start and not context.get('startPeriodDate'):
            context['startPeriodDate'] = derived_start
        if derived_end and not context.get('endPeriodDate'):
            context['endPeriodDate'] = derived_end

    if contract:
        context['rateType'] = contract.rate_type or ''
    _merge_template_variables(context, _collect_contract_variables(contract))

    if client:
        _merge_template_variables(context, _collect_client_variables(client), overwrite=True)
        context.setdefault('companyName', client.name or '')
        context.setdefault('seoFullName', client.signatory or '')
        context.setdefault('companyInn', client.inn or '')
        context.setdefault('companyKpp', client.kpp or '')
        context.setdefault('clientInn', context.get('companyInn', ''))
        context.setdefault('clientKpp', context.get('companyKpp', ''))

    performer_type = payload.performerType or getattr(work_package, 'performer_type', None)

    if contractor:
        _merge_template_variables(context, _collect_contractor_variables(contractor, performer_type))
        context.setdefault('employeeName', contractor.name or '')
        context.setdefault('employeeInn', contractor.inn or '')
        context.setdefault('employeeAddress', contractor.address or '')

    if payload.performerType:
        context['performerType'] = payload.performerType
    elif performer_type:
        context['performerType'] = performer_type

    # Реквизиты договора
    if contract:
        context.setdefault('contractNumber', contract.number or '')
        context.setdefault('employeeContractNumber', contract.number or '')
        # если в ORM есть дата договора — подставь; иначе оставим пустым
        contract_date = getattr(contract, "contract_date", None) or getattr(contract, "created_at", None)
        context['employeeContractDate'] = contract_date.strftime('%d.%m.%Y') if contract_date else ''

    # Короткое ФИО для подписи (если есть отдельное поле — используй его)
    seo_full = context.get('seoFullName') or (client.signatory if client else '')
    context['seoShortName'] = _short_name(seo_full)

    if not context.get('responsiblePerson') and contractor and contractor.name:
        context['responsiblePerson'] = contractor.name

    if not context.get('assignmentBasis') and context.get('employeeContractNumber'):
        basis_parts = [f"Договор № {context['employeeContractNumber']}"]
        if context.get('employeeContractDate'):
            basis_parts[0] += f" от {context['employeeContractDate']}"
        context['assignmentBasis'] = basis_parts[0]

    if not context.get('appendixName') and context.get('employeeContractNumber'):
        context['appendixName'] = f"Приложение к договору № {context['employeeContractNumber']}"

    if not context.get('deadlineDate') and context.get('endPeriodDate'):
        context['deadlineDate'] = context['endPeriodDate']

    context.setdefault('projectSystemName', 'Jira')

    generated_act_number = _derive_act_number(
        context=context,
        work_package=work_package,
        contract=contract,
        payload=payload,
    )
    if generated_act_number:
        context['actNumber'] = generated_act_number

    return context



def _build_default_document(context: dict[str, str]) -> str:
    # Если уже есть сгенерированный связный текст (gptBody) — используем его напрямую
    gpt = (context.get('gptBody') or '').strip()
    if gpt:
        return gpt
    lines = [
        context.get('tableTasks', '')
    ]
    return "\n".join(lines)

# --- Deterministic narrative for act (paragraphs only) ---
def _deterministic_act_html(bullets: list[dict], *, lang: str) -> str:
    """Фоллбэк без GPT: формируем абзацы строго из описаний задач.
    Никаких таблиц, заголовков и списков. Никаких часов/статусов/исполнителей.
    """
    if not bullets:
        return ""
    intro = "During the reporting period the following work was completed:" if lang == "en" else "В отчетный период были выполнены следующие работы:"

    def _strip_html(s: str) -> str:
        s = s or ""
        s = re.sub(r"&lt;", "<", s)
        s = re.sub(r"&gt;", ">", s)
        s = re.sub(r"&amp;", "&", s)
        s = re.sub(r"<[^>]+>", " ", s)         # вычищаем теги в описании
        s = re.sub(r"\s+", " ", s).strip()
        return s

    def _short(s: str, n: int = 1200) -> str:
        s = (s or "").strip()
        return (s[:n] + "…") if len(s) > n else s

    items: list[str] = []
    for b in bullets:
        key = escape(b.get("key", "") or "")
        desc = _strip_html(b.get("description", "") or "")
        if not desc:
            # аккуратный фоллбэк: берём summary, если описания нет
            desc = (b.get("summary", "") or "").strip()
        if not desc:
            continue
        items.append(f"<p><strong>{key}.</strong> {escape(_short(desc))}</p>")

    if not items:
        return ""

    html = [
        '<div class="doc-template doc-template--act">',
        f"<p>{intro}</p>",
        *items,
        "</div>",
    ]
    return "\n".join(html)

def _generate_gpt_act_text(items: Iterable, payload: DocumentCreateRequest) -> str:
    lang = (
        payload.gptOptions.language
        if getattr(payload, "gptOptions", None) and payload.gptOptions and payload.gptOptions.language
        else "ru"
    )

    bullets = _prepare_task_bullets(items)
    plain_html = _build_plain_act_html(bullets, lang=lang)

    gpt_options = getattr(payload, "gptOptions", None)
    if not (_openai_client and gpt_options and getattr(gpt_options, "enabled", False)):
        raise ServiceError(
            "gpt_disabled",
            "Не удалось сформировать связный текст акта: GPT отключён или недоступен.",
        )

    candidate_bullets = [b for b in bullets if _strip_meta_lines(_strip_html_markup(b.get("description") or ""))]
    if not candidate_bullets:
        candidate_bullets = bullets
    if not candidate_bullets:
        return plain_html

    style_map = {
        "neutral": "Нейтральный деловой тон",
        "formal": "Официальный деловой стиль",
        "concise": "Кратко: 1-2 предложения",
        "detailed": "Подробнее: 2-3 предложения",
    }
    if lang == "en":
        style_map = {
            "neutral": "Neutral business tone",
            "formal": "Formal business wording",
            "concise": "Concise, 1-2 sentences",
            "detailed": "More detailed, 2-3 sentences",
        }
    style_hint = style_map.get(getattr(gpt_options, "style", "neutral"), style_map["neutral"])

    extra_notes = getattr(gpt_options, "extraNotes", None) or ""
    period = getattr(payload, "period", "") or (
        f"{getattr(payload, 'startPeriodDate', '')} — {getattr(payload, 'endPeriodDate', '')}"
        if getattr(payload, "startPeriodDate", None) and getattr(payload, "endPeriodDate", None)
        else ""
    )
    doc_type = getattr(payload, "documentType", "Акт") or "Акт"

    task_lines: list[str] = []
    for idx, bullet in enumerate(candidate_bullets[:30], start=1):
        key = bullet.get("key") or ""
        summary = (bullet.get("summary") or "").strip()
        raw_description = bullet.get("description") or ""
        cleaned_description = _strip_meta_lines(_strip_html_markup(raw_description))
        text = cleaned_description or summary
        if not text:
            continue
        hours = bullet.get("hours")
        line = f"{idx}. {key} — {text}"
        if hours:
            line += f" (затрачено {hours} ч.)"
        task_lines.append(line)

    if not task_lines:
        return plain_html

    system_prompt = (
        "Ты — ассистент, который помогает составлять связный текст акта выполненных работ. "
        "Нужно написать несколько предложений, отражающих суть выполненных задач. "
        "Не используй списки, заголовки и таблицы — только абзацы текста."
    )
    if lang == "en":
        system_prompt = (
            "You are an assistant that writes cohesive summaries for statements of work. "
            "Produce a short narrative describing the completed tasks. Use plain paragraphs only, no lists."
        )

    user_instructions = [
        f"Документ: {doc_type}.",
        f"Период: {period}." if period else "",
        f"Стиль: {style_hint}.",
        "Сформируй 2-4 предложения, которые описывают результат работы по задачам.",
        "Укажи ключевые действия, результаты и, если уместно, даты или цифры из описания.",
        "Не перечисляй задачи по одной; объединяй информацию в связный текст.",
        "Не используй HTML и маркеры списков.",
    ]
    if extra_notes:
        user_instructions.append(f"Дополнительные пожелания: {extra_notes}")

    user_prompt = "\n".join(filter(None, user_instructions))
    task_section_header = "Данные задач:" if lang == "ru" else "Task data:"
    task_block = "\n".join(task_lines)

    try:
        response = _openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": f"{user_prompt}\n\n{task_section_header}\n{task_block}",
                },
            ],
            max_tokens=600,
            temperature=0.7,
        )
        generated_text = (
            response.choices[0].message.content if response.choices else ""
        ) or ""
    except Exception as exc:  # pragma: no cover - network/service failures
        logger.warning("GPT generation failed: %s", exc)
        return plain_html

    generated_text = generated_text.strip()
    if not generated_text:
        return plain_html

    raw_paragraphs = re.split(r"\n{2,}", generated_text)
    final_paragraphs: list[str] = []
    for paragraph in raw_paragraphs:
        clean = paragraph.strip().replace("\n", " ")
        clean = re.sub(r"^[\d\-•\)\.(\s]+", "", clean)
        clean = clean.strip()
        if clean:
            final_paragraphs.append(clean)

    if not final_paragraphs:
        return plain_html

    html = ["<div class=\"doc-template doc-template--act\">"]
    html.extend(f"<p>{escape(paragraph)}</p>" for paragraph in final_paragraphs)
    html.append("</div>")
    return "\n".join(html)


def _collect_task_sentences(bullets: list[dict[str, object]], limit: int = 6) -> list[str]:
    sentences: list[str] = []
    for bullet in bullets:
        raw_description = bullet.get("description") or bullet.get("summary") or ""
        cleaned = _strip_meta_lines(_strip_html_markup(str(raw_description)))
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if not cleaned:
            continue
        sentences.append(cleaned)
        if len(sentences) >= limit:
            break
    return sentences


def _build_assignment_fallback(bullets: list[dict[str, object]], context: dict[str, str]) -> dict[str, str]:
    project = context.get('projectName') or context.get('projectKey') or 'проекта'
    start = context.get('startPeriodDate')
    end = context.get('endPeriodDate')
    period_fragment = ''
    if start and end:
        period_fragment = f" за период {start} — {end}"
    total_hours = context.get('totalHours')
    hours_fragment = f", объём работ {total_hours} ч." if total_hours else ''
    company = context.get('companyName')
    customer_phrase = f" заказчику «{company}»" if company else " заказчику"

    sentences = _collect_task_sentences(bullets, limit=5)
    tasks_text = "; ".join(sentences) if sentences else "перечень задач согласно договору"

    order_refs: list[str] = []
    for number_key, date_key in (("orderNumber", "orderDate"), ("orderNumber1", "orderDate1"), ("orderNumber2", "orderDate2")):
        number = context.get(number_key)
        date = context.get(date_key)
        if not number:
            continue
        clause = f"Приказ № {number}"
        if date:
            clause += f" от {date}"
        order_refs.append(clause)
    basis_fragment = "; ".join(order_refs)

    fallback = {
        'assignmentGoal': f"Выполнить задачи проекта {project}{period_fragment}{hours_fragment}.",
        'assignmentPurpose': f"Передать результаты разработки{customer_phrase}.",
        'assignmentRequirements': f"Выполнить следующие работы: {tasks_text}.",
        'assignmentAppendix': "Приложение содержит детализацию задач и артефактов, выполненных в отчётный период.",
        'assignmentBasis': basis_fragment or context.get('assignmentBasis') or "",
    }
    return fallback


def _generate_service_assignment_sections(
    *,
    bullets: list[dict[str, object]],
    context: dict[str, str],
    payload: DocumentCreateRequest,
) -> dict[str, str]:
    if not bullets:
        return {}

    fallback = _build_assignment_fallback(bullets, context)

    if not _openai_client:
        return fallback

    gpt_options = getattr(payload, "gptOptions", None)
    language = getattr(gpt_options, "language", "ru") if gpt_options else "ru"

    def _clean_text(value: object, *, limit: int | None = None) -> str:
        cleaned = _strip_meta_lines(_strip_html_markup(str(value or "")))
        cleaned = re.sub(r"\s+", " ", cleaned).strip()
        if limit and len(cleaned) > limit:
            truncated = cleaned[:limit].rsplit(" ", 1)[0].strip()
            cleaned = truncated or cleaned[:limit].strip()
        return cleaned

    order_fragments: list[str] = []
    basis_candidate = context.get('assignmentBasis') or fallback.get('assignmentBasis')
    if basis_candidate:
        order_fragments.append(basis_candidate)
    for number_key, date_key in (("orderNumber", "orderDate"), ("orderNumber1", "orderDate1"), ("orderNumber2", "orderDate2")):
        number = context.get(number_key)
        if not number:
            continue
        clause = f"Приказ № {number}"
        date = context.get(date_key)
        if date:
            clause += f" от {date}"
        order_fragments.append(clause)
    contract_number = context.get('employeeContractNumber') or context.get('contractNumber')
    contract_date = context.get('employeeContractDate') or context.get('contractDate')
    if contract_number:
        clause = f"Договор № {contract_number}"
        if contract_date:
            clause += f" от {contract_date}"
        order_fragments.append(clause)
    period_start = context.get('startPeriodDate')
    period_end = context.get('endPeriodDate')
    if period_start and period_end:
        order_fragments.append(f"Период работ {period_start} — {period_end}")

    order_info = "; ".join(dict.fromkeys(fragment for fragment in order_fragments if fragment))
    if not order_info:
        order_info = "—"

    narrative_source = context.get('bodygpt') or context.get('gptBody') or ""
    narrative_text = _clean_text(narrative_source, limit=800)
    if not narrative_text:
        sentences = _collect_task_sentences(bullets, limit=5)
        narrative_text = " ".join(sentences)
    if not narrative_text:
        narrative_text = "—"

    info_lines: list[str] = []
    project_label = context.get('projectName') or context.get('projectKey')
    if project_label:
        info_lines.append(f"Проект: {project_label}")
    company_label = context.get('companyName')
    if company_label:
        info_lines.append(f"Заказчик: {company_label}")
    performer_label = context.get('employeeName')
    if performer_label:
        info_lines.append(f"Исполнитель: {performer_label}")
    period_start = context.get('startPeriodDate')
    period_end = context.get('endPeriodDate')
    if period_start and period_end:
        info_lines.append(f"Период: {period_start} — {period_end}")
    hours_label = context.get('totalHours')
    if hours_label:
        info_lines.append(f"Объём часов: {hours_label}")
    amount_label = context.get('totalAmount')
    if amount_label:
        info_lines.append(f"Сумма: {amount_label}")
    software_label = context.get('softwareName') or context.get('projectName')
    if software_label:
        info_lines.append(f"ПО: {software_label}")

    tasks_lines: list[str] = []
    for index, bullet in enumerate(bullets[:20], start=1):
        key = str(bullet.get('key') or index)
        description = _strip_meta_lines(_strip_html_markup(str(bullet.get('description') or '')))
        if not description:
            description = str(bullet.get('summary') or '').strip()
        description = re.sub(r"\s+", " ", description).strip()
        if not description:
            continue
        hours = bullet.get('hours')
        hour_fragment = f" (затрачено {hours} ч.)" if hours else ""
        tasks_lines.append(f"{index}. {key}: {description}{hour_fragment}")

    if not tasks_lines:
        tasks_lines.append("—")

    if language == "en":
        system_prompt = (
            "You prepare service assignment sections based strictly on the provided order and completed tasks. "
            "Use only the supplied facts. Output must be valid JSON."
        )
        user_head = (
            "Create a JSON object with keys goal, basis, purpose, requirements, appendix.\n"
            "Each value must be one concise sentence in a formal business tone. No lists or numbering. Use '—' if data is missing.\n"
            "Combine the development order, act narrative, and task breakdown."
        )
        instructions_block = (
            "- goal: describe the main development objective, referencing the software or module.\n"
            "- basis: if the order/basis string equals '—', craft a lawful basis using the contract and period; otherwise repeat the provided string verbatim.\n"
            "- purpose: state the business outcome for the customer or end users.\n"
            "- requirements: summarise the essential functional changes from the tasks in one sentence.\n"
            "- appendix: mention that the appendix contains the detailed list of tasks, hours, and amount if available."
        )
        order_label = "Development order / basis:"
        narrative_label = "Narrative from the act:"
        tasks_label = "Task breakdown:"
        context_label = "Additional context:"
    else:
        system_prompt = (
            "Ты готовишь разделы служебного задания, опираясь на приказ на разработку и фактические задачи. "
            "Используй только переданные факты, не придумывай новые данные. Ответ возвращай строго в формате JSON."
        )
        user_head = (
            "Сформируй JSON с ключами goal, basis, purpose, requirements, appendix.\n"
            "Каждое значение — одно официально-деловое предложение без перечней и нумерации. Если данных нет, используй символ '—'.\n"
            "Необходимо объединить информацию из приказа, акта и списка задач."
        )
        instructions_block = (
            "- goal: обозначь ключевую цель разработки, упоминая продукт или модуль.\n"
            "- basis: если строка «Приказ / основание» равна «—», сформулируй правовое основание с использованием договора и периода; иначе повтори её дословно.\n"
            "- purpose: укажи, какой результат получает заказчик или пользователи.\n"
            "- requirements: через одно предложение перечисли основные доработки и функции из задач.\n"
            "- appendix: подчеркни, что приложение содержит детализацию задач, часы и сумму (если данные есть)."
        )
        order_label = "Приказ / основание:"
        narrative_label = "Описание выполненных работ из акта:"
        tasks_label = "Детализация задач:"
        context_label = "Контекст:"

    info_block = "\n".join(info_lines)
    tasks_block = "\n".join(tasks_lines)

    prompt_sections = [
        user_head,
        instructions_block,
        f"{order_label}\n{order_info}",
        f"{narrative_label}\n{narrative_text}",
        f"{tasks_label}\n{tasks_block}",
    ]
    if info_block:
        prompt_sections.append(f"{context_label}\n{info_block}")
    user_prompt = "\n\n".join(prompt_sections)

    try:
        response = _openai_client.chat.completions.create(
            model=OPENAI_MODEL,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            max_tokens=400,
            temperature=0,
            response_format={"type": "json_object"},
        )
        content = response.choices[0].message.content if response.choices else ""
        data = json.loads(content) if content else {}
    except Exception as exc:  # pragma: no cover - network/format failures
        logger.debug("GPT assignment generation failed: %s", exc)
        return fallback

    mapping = {
        'goal': 'assignmentGoal',
        'basis': 'assignmentBasis',
        'purpose': 'assignmentPurpose',
        'requirements': 'assignmentRequirements',
        'appendix': 'assignmentAppendix',
    }
    result: dict[str, str] = {}
    for source_key, target_key in mapping.items():
        value = data.get(source_key)
        if isinstance(value, list):
            value = " ".join(str(item).strip() for item in value if str(item).strip())
        if isinstance(value, str):
            candidate = value.strip()
            if candidate:
                result[target_key] = candidate

    for key, fallback_value in fallback.items():
        result.setdefault(key, fallback_value)

    if order_info and order_info != "—":
        result['assignmentBasis'] = order_info

    if not result.get('assignmentBasis'):
        basis_values = [fallback.get('assignmentBasis'), context.get('assignmentBasis')]
        for value in basis_values:
            if value:
                result['assignmentBasis'] = value
                break
        else:
            orders = []
            for number_key, date_key in (("orderNumber", "orderDate"), ("orderNumber1", "orderDate1"), ("orderNumber2", "orderDate2")):
                num = context.get(number_key)
                date = context.get(date_key)
                if not num:
                    continue
                clause = f"Приказ № {num}"
                if date:
                    clause += f" от {date}"
                orders.append(clause)
            if orders:
                result['assignmentBasis'] = "; ".join(orders)
            else:
                contract_number = context.get('employeeContractNumber') or context.get('contractNumber')
                contract_date = context.get('employeeContractDate') or context.get('contractDate')
                if contract_number:
                    clause = f"Договор № {contract_number}"
                    if contract_date:
                        clause += f" от {contract_date}"
                    result['assignmentBasis'] = clause
                elif period_start and period_end:
                    result['assignmentBasis'] = f"Работы по периоду {period_start} — {period_end}"

    return result
def _boolish(v) -> bool:
    if isinstance(v, bool):
        return v
    if isinstance(v, (int, float)):
        return v != 0
    if isinstance(v, str):
        return v.strip().lower() in {"1", "true", "yes", "y", "on"}
    return False

def _has_description(obj) -> bool:
    try:
        desc = getattr(obj, "description", None)
    except Exception:
        desc = None
    if desc is None and isinstance(obj, dict):
        desc = obj.get("description")
    return bool(desc and str(desc).strip())

def _issue_key(identifier: str) -> str:
    if ":" in identifier:
        return identifier.split(":", 1)[1]
    return identifier


def _map_metadata(raw: dict | None) -> WorkPackageMetadata:
    payload = raw or {}
    prepared_for = [
        str(value).strip()
        for value in payload.get("preparedFor", [])
        if isinstance(value, str) and str(value).strip()
    ]
    tags = [
        str(value).strip()
        for value in payload.get("tags", [])
        if isinstance(value, str) and str(value).strip()
    ]
    unique_tags = list(dict.fromkeys(tags))

    return WorkPackageMetadata(
        preparedFor=prepared_for,
        tags=unique_tags,
        taxCategory=payload.get("taxCategory"),
        benefitCategory=payload.get("benefitCategory"),
        currency=payload.get("currency", "RUB"),
        adjustments=payload.get("adjustments"),
    )


def _map_task_snapshots(raw: Iterable[dict]) -> list[TaskCostSnapshot]:
    return [TaskCostSnapshot(**snapshot) for snapshot in raw]


def _map_work_package(record: orm_models.WorkPackageORM) -> WorkPackage:
    return WorkPackage(
        id=record.id,
        createdAt=record.created_at,
        period=record.period,
        projectKey=record.project_key,
        projectName=record.project_name,
        contractId=record.contract_id,
        clientId=record.client_id,
        contractorId=record.contractor_id,
        totalHours=record.total_hours,
        totalAmount=record.total_amount,
        hourlyRate=record.hourly_rate,
        baseRate=record.base_rate,
        rateType=record.rate_type,
        includeTimesheet=record.include_timesheet,
        currency=record.currency,
        performerType=record.performer_type,
        vatIncluded=record.vat_included,
        vatPercent=record.vat_percent,
        vatAmount=record.vat_amount,
        taskSnapshots=_map_task_snapshots(record.task_snapshots or []),
        metadata=_map_metadata(record.metadata_json),
    )


def _virtual_work_package(work_package_id: str) -> WorkPackage:
    return WorkPackage(
        id=work_package_id,
        createdAt=datetime.utcnow(),
        period="",
        projectKey="",
        projectName="",
        contractId="",
        clientId="",
        contractorId="",
        totalHours=0.0,
        totalAmount=0.0,
        hourlyRate=0.0,
        baseRate=0.0,
        rateType="hour",
        includeTimesheet=False,
        currency="RUB",
        performerType="individual",
        vatIncluded=False,
        vatPercent=0.0,
        vatAmount=0.0,
        taskSnapshots=[],
        metadata=WorkPackageMetadata(),
    )


def _doc_v2_priority(record: orm_models.DocumentV2ORM) -> int:
    return DOC_V2_PRIORITY.get(record.doc_type or "", 100)


def _format_period_range(start: date | None, end: date | None, fallback: str) -> str:
    start_text = start.isoformat() if start else ""
    end_text = end.isoformat() if end else ""
    if start_text and end_text:
        return f"{start_text} — {end_text}"
    return start_text or end_text or fallback


def _map_v2_document_to_work_package(
    session: Session,
    record: orm_models.DocumentV2ORM,
    performer_index: int,
) -> WorkPackage:
    doc_record = _map_closing_document(session, record)
    period = _format_period_range(record.period_start, record.period_end, doc_record.period)
    metadata = doc_record.metadata
    currency = metadata.currency if metadata and getattr(metadata, "currency", None) else "RUB"
    performer_id = doc_record.contractorId or None

    return WorkPackage(
        id=f"{doc_record.workPackageId}-perf-{performer_index}",
        createdAt=doc_record.createdAt,
        period=period,
        projectKey=doc_record.projectKey,
        projectName=doc_record.projectName,
        contractId=doc_record.contractId,
        clientId=doc_record.clientId,
        contractorId=doc_record.contractorId,
        totalHours=doc_record.totalHours,
        totalAmount=doc_record.amount,
        hourlyRate=doc_record.hourlyRate,
        baseRate=doc_record.baseRate,
        rateType=doc_record.rateType,
        includeTimesheet=doc_record.includeTimesheet,
        currency=currency,
        performerType=doc_record.performerType,
        vatIncluded=doc_record.vatIncluded,
        vatPercent=doc_record.vatPercent,
        vatAmount=doc_record.vatAmount,
        taskSnapshots=doc_record.taskSnapshots,
        metadata=metadata,
        performerId=performer_id,
    )


def release_work_package_tasks(session: Session, work_package_id: str) -> WorkPackage:
    def _release_tasks_by_prefix(prefix: str) -> None:
        tasks = (
            session.execute(
                select(orm_models.TaskORM).where(orm_models.TaskORM.work_package_id.like(f"{prefix}%"))
            )
            .scalars()
            .all()
        )
        for task in tasks:
            task.work_package_id = None
            task.force_included = False
        if tasks:
            session.flush()

    if work_package_id.startswith("package-v2-"):
        suffix = work_package_id.removeprefix("package-v2-")
        if not suffix:
            raise ValueError("Work package not found")
        _release_tasks_by_prefix(f"package:{suffix}")
        return _virtual_work_package(work_package_id)

    if work_package_id.startswith("package:"):
        _release_tasks_by_prefix(work_package_id)
        return _virtual_work_package(work_package_id)

    work_package = session.get(orm_models.WorkPackageORM, work_package_id)
    if not work_package:
        raise ValueError("Work package not found")

    _release_tasks_by_prefix(work_package_id)
    return _map_work_package(work_package)


def _map_assignee(user: orm_models.UserORM | None) -> DocumentAssignee | None:
    if not user:
        return None
    full_name = (user.full_name or "").strip()
    return DocumentAssignee(
        id=user.id,
        email=user.email,
        full_name=full_name or user.email,
    )


def _map_document(record: orm_models.DocumentRecordORM) -> DocumentRecord:
    work_package = record.work_package
    snapshots_source = work_package.task_snapshots if work_package else record.metadata_json.get("taskSnapshots", [])
    status = (record.approval_status or "draft").strip()
    if status == "performer_approved":
        status = "pending_manager"
    elif status not in {
        "draft",
        "pending_performer",
        "pending_manager",
        "manager_approved",
        "rejected_performer",
        "rejected_manager",
        "final",
    }:
        # fallback for legacy values
        status = "draft"

    performer_assignee = _map_assignee(record.performer_assignee)
    manager_assignee = _map_assignee(record.manager_assignee)

    workspace_name = None
    if record.workspace is not None:
        workspace_name = record.workspace.name

    return DocumentRecord(
        id=record.id,
        createdAt=record.created_at,
        period=record.period,
        type=record.type,
        workspace_id=record.workspace_id,
        workspace_name=workspace_name,
        clientId=record.client_id,
        contractorId=record.contractor_id,
        contractId=record.contract_id,
        projectKey=record.project_key,
        projectName=record.project_name,
        tasksCount=record.task_count,
        totalHours=record.total_hours,
        amount=record.total_amount,
        hourlyRate=record.hourly_rate,
        baseRate=record.base_rate,
        rateType=record.rate_type,
        status=record.status,
        includeTimesheet=record.include_timesheet,
        files=[DocumentFile(**file_meta) for file_meta in record.files or []],
        taskSnapshots=_map_task_snapshots(snapshots_source),
        workPackageId=record.work_package_id,
        notes=record.notes or None,
        templateId=record.template_id,
        performerType=record.performer_type,
        vatIncluded=record.vat_included,
        vatPercent=record.vat_percent,
        vatAmount=record.vat_amount,
        metadata=_map_metadata(record.metadata_json),
        approvalStatus=status,
        submittedAt=record.submitted_at,
        managerApprovedAt=record.manager_approved_at,
        managerApprovedBy=record.manager_approved_by,
        performerApprovedAt=record.performer_approved_at,
        performerApprovedBy=record.performer_approved_by,
        finalizedAt=record.finalized_at,
        finalizedBy=record.finalized_by,
        approvalNotes=[
            DocumentApprovalNote(
                timestamp=note.get("timestamp"),
                author=note.get("author", ""),
                role=note.get("role", ""),
                status=note.get("status", ""),
                message=note.get("message", ""),
            )
            if isinstance(note, dict)
            else note
            for note in (record.approval_notes or [])
        ],
        performerAssignee=performer_assignee,
        managerAssignee=manager_assignee,
        shared_with_parent=bool(record.shared_with_parent),
        shared_parent_id=record.shared_parent_id,
        shared_at=record.shared_at,
        shared_by_user_id=record.shared_by_user_id,
    )


def _hydrate_document_assignees(session: Session, record: orm_models.DocumentRecordORM) -> None:
    performer_missing = not record.performer_assignee_id
    manager_missing = not record.manager_assignee_id
    if not performer_missing and not manager_missing:
        return

    contractor = session.get(orm_models.IndividualORM, record.contractor_id)
    if not contractor:
        return

    updated = False

    performer_user, _ = ensure_individual_account(
        session,
        contractor,
        create_missing=False,
        allow_email_reassign=True,
    )

    if performer_missing and performer_user:
        record.performer_assignee = performer_user
        record.performer_assignee_id = performer_user.id
        updated = True

    if manager_missing and contractor.default_manager_id:
        manager = session.get(orm_models.IndividualORM, contractor.default_manager_id)
        if manager is not None:
            manager_user, _ = ensure_individual_account(
                session,
                manager,
                create_missing=False,
                allow_email_reassign=True,
            )
            if manager_user is not None:
                record.manager_assignee = manager_user
                record.manager_assignee_id = manager_user.id
                updated = True

    if updated:
        session.flush()


def _safe_float(value: object, *, precision: int | None = None) -> float:
    if value is None:
        return 0.0
    if isinstance(value, Decimal):
        result = float(value)
    else:
        try:
            result = float(value)  # type: ignore[arg-type]
        except (TypeError, ValueError):
            return 0.0
    if precision is not None:
        return float(round(result, precision))
    return result


def _build_task_snapshots_from_items(
    items: List[dict],
    *,
    fallback_rate: float,
    project_key: str,
    project_name: str,
) -> List[TaskCostSnapshot]:
    snapshots: List[TaskCostSnapshot] = []
    for raw in items:
        if not isinstance(raw, dict):
            continue
        hours = _safe_float(raw.get("hours"), precision=4)
        rate = _safe_float(raw.get("hourlyRate"), precision=4)
        if rate <= 0 and fallback_rate > 0:
            rate = fallback_rate
        amount = _safe_float(raw.get("amount"), precision=2)
        if amount <= 0 and rate > 0 and hours > 0:
            amount = round(rate * hours, 2)

        snapshot = TaskCostSnapshot(
            id=str(raw.get("id") or raw.get("jira_id") or uuid4()),
            key=str(raw.get("key") or raw.get("jira_id") or "UNKNOWN"),
            title=str(raw.get("summary") or ""),
            description=str(raw.get("description")) if isinstance(raw.get("description"), str) else None,
            status=str(raw.get("status") or ""),
            hours=round(hours, 2),
            billable=bool(raw.get("billable", True)),
            forceIncluded=bool(raw.get("forceIncluded") or raw.get("force_included") or False),
            projectKey=str(raw.get("projectKey") or raw.get("project_key") or project_key),
            projectName=str(raw.get("projectName") or raw.get("project_name") or project_name),
            hourlyRate=round(rate, 2),
            amount=round(amount, 2),
            categories=None,
        )
        snapshots.append(snapshot)
    return snapshots


def _map_closing_document(session: Session, record: orm_models.DocumentV2ORM) -> DocumentRecord:
    meta = record.meta if isinstance(record.meta, dict) else {}
    package = record.package
    contract_v2 = record.contract
    performer_v2 = record.performer or (contract_v2.performer if contract_v2 else None)
    company = contract_v2.company if contract_v2 else None

    approval_meta = meta.get("approval") if isinstance(meta, dict) else None
    approval_status = "draft"
    submitted_at = None
    manager_approved_at = None
    manager_approved_by = None
    performer_approved_at = None
    performer_approved_by = None
    finalized_at = None
    finalized_by = None
    approval_notes: List[DocumentApprovalNote] = []

    performer_assignee_obj: DocumentAssignee | None = None
    manager_assignee_obj: DocumentAssignee | None = None

    def _parse_dt(value: object) -> datetime | None:
        if isinstance(value, datetime):
            return value
        if isinstance(value, str):
            text = value.strip()
            if not text:
                return None
            try:
                return datetime.fromisoformat(text.replace("Z", "+00:00"))
            except ValueError:
                return None
        if isinstance(value, (int, float)):
            try:
                return datetime.fromtimestamp(float(value))
            except (ValueError, OSError):
                return None
        return None

    performer_assignee_obj: DocumentAssignee | None = None
    manager_assignee_obj: DocumentAssignee | None = None

    approval_changed = False

    if isinstance(approval_meta, dict):
        approval_status = _normalize_v2_status(approval_meta.get("status"))

        def _parse_assignee(entry: object) -> DocumentAssignee | None:
            if not isinstance(entry, dict):
                return None
            assignee_id = entry.get("id")
            email = entry.get("email")
            full_name = entry.get("full_name") or entry.get("fullName")
            if not isinstance(assignee_id, str) or not assignee_id:
                return None
            if not isinstance(email, str) or not email:
                return None
            if not isinstance(full_name, str) or not full_name.strip():
                full_name = email
            return DocumentAssignee(id=assignee_id, email=email, full_name=full_name)

        performer_assignee_obj = _parse_assignee(approval_meta.get("performer_assignee"))
        manager_assignee_obj = _parse_assignee(approval_meta.get("manager_assignee"))

        def _parse_assignee(entry: object) -> DocumentAssignee | None:
            if not isinstance(entry, dict):
                return None
            assignee_id = entry.get("id")
            email = entry.get("email")
            full_name = entry.get("full_name") or entry.get("fullName")
            if not isinstance(assignee_id, str) or not assignee_id:
                return None
            if not isinstance(email, str) or not email:
                return None
            if not isinstance(full_name, str) or not full_name.strip():
                full_name = email
            return DocumentAssignee(id=assignee_id, email=email, full_name=full_name)

        performer_assignee_obj = _parse_assignee(approval_meta.get("performer_assignee"))
        manager_assignee_obj = _parse_assignee(approval_meta.get("manager_assignee"))

        timeline = approval_meta.get("timeline") or approval_meta.get("history") or approval_meta.get("notes")
        if isinstance(timeline, list):
            for entry in timeline:
                if not isinstance(entry, dict):
                    continue
                status_entry = _normalize_v2_status(entry.get("status") or approval_status)
                timestamp_entry = _parse_dt(
                    entry.get("timestamp")
                    or entry.get("time")
                    or entry.get("created_at")
                    or entry.get("createdAt")
                )
                note_text = entry.get("message") or entry.get("note") or entry.get("comment") or ""
                author = entry.get("author") or entry.get("user") or ""
                role = entry.get("role") or entry.get("actor_role") or ""

                approval_notes.append(
                    DocumentApprovalNote(
                        timestamp=timestamp_entry or datetime.utcnow(),
                        author=str(author or ""),
                        role=str(role or ""),
                        status=status_entry,
                        message=str(note_text or ""),
                    )
                )

        submitted_at = _parse_dt(approval_meta.get("submitted_at") or approval_meta.get("submittedAt"))
        manager_approved_at = _parse_dt(approval_meta.get("manager_approved_at") or approval_meta.get("managerApprovedAt"))
        performer_approved_at = _parse_dt(approval_meta.get("performer_approved_at") or approval_meta.get("performerApprovedAt"))
        finalized_at = _parse_dt(approval_meta.get("finalized_at") or approval_meta.get("finalizedAt"))

        manager_approved_by = approval_meta.get("manager_approved_by") or approval_meta.get("managerApprovedBy")
        performer_approved_by = approval_meta.get("performer_approved_by") or approval_meta.get("performerApprovedBy")
        finalized_by = approval_meta.get("finalized_by") or approval_meta.get("finalizedBy")

    if finalized_at is None and approval_status == "final":
        finalized_at = record.updated_at or record.created_at


    legacy_contract_id = None
    if isinstance(meta, dict):
        legacy_contract_id = meta.get("legacy_contract_id") or meta.get("legacyContractId")
    legacy_contract = (
        session.get(orm_models.ContractORM, str(legacy_contract_id))
        if legacy_contract_id
        else None
    )
    if legacy_contract is not None and object_session(legacy_contract) is None:
        legacy_contract = session.merge(legacy_contract, load=False)
    legacy_performer = None
    if legacy_contract and legacy_contract.contractor_id:
        legacy_performer = session.get(orm_models.IndividualORM, legacy_contract.contractor_id)
        if legacy_performer is not None:
            ensure_individual_account(
                session,
                legacy_performer,
                create_missing=False,
                allow_email_reassign=True,
            )

    if performer_assignee_obj is None and legacy_performer and legacy_performer.user:
        user = legacy_performer.user
        performer_assignee_obj = DocumentAssignee(
            id=user.id,
            email=user.email,
            full_name=user.full_name or user.email,
        )
        approval_changed = True

    if manager_assignee_obj is None and legacy_performer and legacy_performer.default_manager_id:
        manager = session.get(orm_models.IndividualORM, legacy_performer.default_manager_id)
        if manager is not None:
            ensure_individual_account(
                session,
                manager,
                create_missing=False,
                allow_email_reassign=True,
            )
            if manager.user:
                manager_assignee_obj = DocumentAssignee(
                    id=manager.user.id,
                    email=manager.user.email,
                    full_name=manager.user.full_name or manager.user.email,
                )
                approval_changed = True

    def _assignee_to_meta(obj: DocumentAssignee | None) -> dict | None:
        if obj is None:
            return None
        return {
            "id": obj.id,
            "email": obj.email,
            "full_name": obj.fullName or obj.email,
        }

    if isinstance(approval_meta, dict):
        meta_performer = _assignee_to_meta(performer_assignee_obj)
        if meta_performer and approval_meta.get("performer_assignee") != meta_performer:
            approval_meta["performer_assignee"] = meta_performer
            approval_changed = True
        meta_manager = _assignee_to_meta(manager_assignee_obj)
        if meta_manager and approval_meta.get("manager_assignee") != meta_manager:
            approval_meta["manager_assignee"] = meta_manager
            approval_changed = True

    if approval_changed:
        meta["approval"] = approval_meta
        record.meta = meta
        flag_modified(record, "meta")
        session.flush()

    project_key = str(meta.get("project_key") or meta.get("projectKey") or "")
    project_name = str(meta.get("project_name") or meta.get("projectName") or project_key)

    task_items_raw: List[dict] = []
    raw_items = meta.get("task_items") or meta.get("taskItems")
    if isinstance(raw_items, list):
        task_items_raw = [item for item in raw_items if isinstance(item, dict)]
    if not task_items_raw and record.timesheet and isinstance(record.timesheet.task_table, list):
        for entry in record.timesheet.task_table:
            if not isinstance(entry, dict):
                continue
            task_items_raw.append(
                {
                    "id": entry.get("jira_id") or entry.get("id"),
                    "jira_id": entry.get("jira_id") or entry.get("id"),
                    "key": entry.get("jira_id") or entry.get("id"),
                    "summary": entry.get("summary") or "",
                    "status": entry.get("status") or "",
                    "hours": entry.get("hours"),
                    "billable": entry.get("billable", True),
                }
            )

    raw_hours_total = sum(_safe_float(item.get("hours")) for item in task_items_raw)
    hours_total = _safe_float(record.hours)
    if hours_total <= 0 and raw_hours_total > 0:
        hours_total = raw_hours_total

    amount_total = _safe_float(record.amount_total)
    if amount_total <= 0:
        amount_total = sum(_safe_float(item.get("amount")) for item in task_items_raw)

    rate_hour = _safe_float(record.rate_hour, precision=4)
    if rate_hour <= 0 and hours_total > 0 and amount_total > 0:
        rate_hour = round(amount_total / hours_total, 4)

    task_snapshots = _build_task_snapshots_from_items(
        task_items_raw,
        fallback_rate=rate_hour,
        project_key=project_key,
        project_name=project_name,
    )

    tasks_count = len(task_snapshots) or len(task_items_raw)
    if tasks_count == 0:
        tasks_raw = meta.get("tasks")
        if isinstance(tasks_raw, list):
            tasks_count = len(tasks_raw)

    vat_amount = _safe_float(record.vat_amount, precision=2)
    amount_wo_vat = _safe_float(record.amount_wo_vat, precision=2)
    vat_percent = _safe_float(meta.get("vat_percent") or meta.get("vatPercent"), precision=2)
    if vat_percent <= 0 and amount_wo_vat > 0 and vat_amount > 0:
        vat_percent = round((vat_amount / amount_wo_vat) * 100, 2)
    elif vat_percent <= 0 and amount_total > 0 and vat_amount > 0:
        vat_percent = round((vat_amount / (amount_total - vat_amount)) * 100, 2) if amount_total > vat_amount else 0.0

    vat_included = vat_amount > 0

    prepared_for = DOC_V2_AUDIENCE.get(record.doc_type, ["act"])
    tags = [
        tag
        for tag in {
            "package-v2",
            record.doc_type.lower() if record.doc_type else "",
            project_key,
            project_name,
            package.period_start.strftime("%Y-%m") if package and package.period_start else "",
        }
        if tag
    ]

    performer_name = None
    if legacy_contract and legacy_contract.contractor:
        performer_name = legacy_contract.contractor.name
    elif performer_v2 and performer_v2.full_name:
        performer_name = performer_v2.full_name

    client_name = None
    if legacy_contract and legacy_contract.client:
        client_name = legacy_contract.client.name
    elif company and company.name:
        client_name = company.name

    contract_number = None
    if legacy_contract and legacy_contract.number:
        contract_number = legacy_contract.number
    elif contract_v2 and contract_v2.contract_number:
        contract_number = contract_v2.contract_number

    metadata = WorkPackageMetadata(
        preparedFor=prepared_for,
        tags=tags,
        taxCategory=None,
        benefitCategory=None,
        currency=(contract_v2.currency if contract_v2 and contract_v2.currency else "RUB"),
        adjustments=None,
    )

    include_timesheet = bool(meta.get("include_timesheet") or meta.get("includeTimesheet") or record.timesheet)

    doc_type_label = DOC_V2_TYPE_LABELS.get(record.doc_type, "Пакет")
    file_type = DOC_V2_FILE_TYPES.get(record.doc_type, "package")
    file_entry = DocumentFile(
        id=f"package-v2-{record.id}-main",
        label=doc_type_label,
        type=file_type,
        format="docx",
        status="Готов",
        url=f"/packages/{record.package_id}/documents/{record.id}/file",
    )

    contract_id_value: str
    client_id_value: str
    contractor_id_value: str

    if legacy_contract:
        contract_id_value = str(legacy_contract.id)
        client_id_value = str(legacy_contract.client_id or f"package-v2-client-{record.package_id}")
        contractor_id_value = str(legacy_contract.contractor_id or f"package-v2-performer-{record.package_id}")
    else:
        contract_id_value = f"package-v2-contract-{record.contract_id}"
        client_id_value = f"package-v2-client-{contract_v2.company_id if contract_v2 and contract_v2.company_id is not None else record.package_id}"
        contractor_id_value = f"package-v2-performer-{performer_v2.id if performer_v2 and performer_v2.id is not None else record.package_id}"

    period_value = None
    if record.period_start:
        period_value = record.period_start.strftime("%Y-%m")
    elif package and package.period_start:
        period_value = package.period_start.strftime("%Y-%m")
    else:
        period_value = record.created_at.strftime("%Y-%m") if record.created_at else ""

    performer_type = str(
        meta.get("performer_type")
        or (contract_v2.performer.type if contract_v2 and contract_v2.performer and contract_v2.performer.type else None)
        or "gph"
    )

    def _extract_workspace_id(source: object) -> str | None:
        if isinstance(source, str):
            candidate = source.strip()
            return candidate or None
        if isinstance(source, dict):
            candidate: object = (
                source.get("workspace_id")
                or source.get("workspaceId")
                or source.get("workspace")
            )
            if isinstance(candidate, dict):
                inner = candidate.get("id") or candidate.get("workspace_id") or candidate.get("workspaceId")
                candidate = inner
            if isinstance(candidate, str):
                candidate = candidate.strip()
                if candidate:
                    return candidate
        return None

    workspace_id = _extract_workspace_id(meta)
    if workspace_id is None and package is not None:
        workspace_id = _extract_workspace_id(getattr(package, "meta", None))
    if workspace_id is None and contract_v2 is not None:
        workspace_id = _extract_workspace_id(getattr(contract_v2, "meta", None))
    if workspace_id is None:
        workspace_id = session.info.get("workspace_id")
    if not workspace_id:
        raise ValueError("Не удалось определить рабочее пространство для документа")

    workspace_name = None
    workspace = session.get(orm_models.WorkspaceORM, workspace_id)
    if workspace is not None:
        workspace_name = workspace.name

    return DocumentRecord(
        id=f"package-v2-{record.id}",
        createdAt=record.created_at or datetime.utcnow(),
        period=period_value,
        type=doc_type_label,
        workspace_id=workspace_id,
        workspace_name=workspace_name,
        clientId=client_id_value,
        contractorId=contractor_id_value,
        contractId=contract_id_value,
        projectKey=project_key,
        projectName=project_name,
        tasksCount=tasks_count,
        totalHours=round(hours_total, 2),
        amount=round(amount_total, 2),
        hourlyRate=round(rate_hour, 2),
        baseRate=round(rate_hour, 2),
        rateType="hour",
        status="Не согласован",
        includeTimesheet=include_timesheet,
        files=[file_entry],
        taskSnapshots=task_snapshots,
        workPackageId=f"package-v2-{record.package_id}",
        notes=None,
        templateId=str(record.template_id) if record.template_id is not None else None,
        performerType=performer_type,
        vatIncluded=vat_included,
        vatPercent=round(vat_percent, 2),
        vatAmount=round(vat_amount, 2),
        metadata=metadata,
        approvalStatus=approval_status,
        submittedAt=submitted_at,
        managerApprovedAt=manager_approved_at,
        managerApprovedBy=str(manager_approved_by) if manager_approved_by is not None else None,
        performerApprovedAt=performer_approved_at,
        performerApprovedBy=str(performer_approved_by) if performer_approved_by is not None else None,
        finalizedAt=finalized_at,
        finalizedBy=str(finalized_by) if finalized_by is not None else None,
        approvalNotes=approval_notes,
        performerAssignee=performer_assignee_obj,
        managerAssignee=manager_assignee_obj,
    )


def _deduplicate_notes(notes: Iterable[DocumentApprovalNote]) -> list[DocumentApprovalNote]:
    seen: set[tuple] = set()
    ordered: list[DocumentApprovalNote] = []
    for note in sorted(
        notes,
        key=lambda item: (item.timestamp or datetime.min, item.status, item.author, item.message),
    ):
        signature = (
            (note.timestamp.isoformat() if note.timestamp else None),
            note.status,
            note.author,
            note.role,
            note.message,
        )
        if signature in seen:
            continue
        seen.add(signature)
        ordered.append(note)
    return ordered


def _merge_closing_documents(
    session: Session,
    records: list[orm_models.DocumentV2ORM],
) -> DocumentRecord:
    if not records:
        raise ValueError("Expected at least one closing document")

    mapped_documents = [_map_closing_document(session, record) for record in records]
    base_document = mapped_documents[0]

    # Aggregate files with deterministic ordering
    file_entries: list[tuple[int, DocumentFile]] = []
    doc_labels: list[str] = []
    for record, mapped in zip(records, mapped_documents):
        priority = DOC_V2_PRIORITY.get(record.doc_type or "", 100)
        for file_entry in mapped.files:
            file_entries.append((priority, file_entry))
        label = DOC_V2_TYPE_LABELS.get(record.doc_type or "", "Документ")
        doc_labels.append(label)
    file_entries.sort(key=lambda item: (item[0], item[1].label))
    merged_files: dict[str, DocumentFile] = {}
    for _, file_entry in file_entries:
        merged_files.setdefault(file_entry.id, file_entry)

    # Aggregate task snapshots
    snapshot_map: dict[str, TaskCostSnapshot] = {}
    for mapped in mapped_documents:
        for snapshot in mapped.taskSnapshots:
            snapshot_map[snapshot.id] = snapshot

    unique_snapshots = list(snapshot_map.values())
    unique_snapshots.sort(key=lambda snapshot: snapshot.id)

    total_hours = round(sum(snapshot.hours for snapshot in unique_snapshots), 2)
    total_amount = round(sum(snapshot.amount for snapshot in unique_snapshots), 2)
    if total_amount <= 0:
        total_amount = max((mapped.amount for mapped in mapped_documents), default=0.0)
    if total_hours <= 0:
        total_hours = max((mapped.totalHours for mapped in mapped_documents), default=0.0)

    vat_amount_candidates = [mapped.vatAmount for mapped in mapped_documents if mapped.vatAmount > 0]
    vat_amount = max(vat_amount_candidates, default=0.0)
    vat_percent_candidates = [mapped.vatPercent for mapped in mapped_documents if mapped.vatPercent > 0]
    vat_percent = max(vat_percent_candidates, default=base_document.vatPercent)
    vat_included = any(mapped.vatIncluded or mapped.vatAmount > 0 for mapped in mapped_documents)

    tags: list[str] = []
    prepared_for: list[str] = []
    for mapped in mapped_documents:
        tags.extend(mapped.metadata.tags or [])
        prepared_for.extend(mapped.metadata.preparedFor or [])
    deduped_tags = list(dict.fromkeys(tags))
    deduped_prepared = list(dict.fromkeys(prepared_for))

    include_timesheet = any(mapped.includeTimesheet for mapped in mapped_documents)

    status_priority = {
        "rejected_performer": 0,
        "rejected_manager": 0,
        "draft": 10,
        "pending_performer": 20,
        "pending_manager": 30,
        "manager_approved": 40,
        "final": 50,
    }
    aggregated_status = base_document.approvalStatus
    aggregated_priority = status_priority.get(aggregated_status, 10)
    collected_notes: list[DocumentApprovalNote] = []

    for mapped in mapped_documents:
        priority = status_priority.get(mapped.approvalStatus, 10)
        if priority < aggregated_priority:
            aggregated_status = mapped.approvalStatus
            aggregated_priority = priority
        collected_notes.extend(mapped.approvalNotes or [])

    aggregated_metadata = base_document.metadata.copy(update={
        "tags": deduped_tags,
        "preparedFor": deduped_prepared,
    })

    hourly_rate = base_document.hourlyRate
    if total_hours > 0 and total_amount > 0:
        hourly_rate = round(total_amount / total_hours, 2)

    doc_labels = list(dict.fromkeys(doc_labels))
    package_label = "Пакет документов"
    if doc_labels:
        package_label = f"Пакет: {', '.join(doc_labels)}"

    aggregated_document = base_document.copy(update={
        "type": package_label,
        "files": list(merged_files.values()),
        "taskSnapshots": unique_snapshots,
        "tasksCount": len(unique_snapshots),
        "totalHours": total_hours,
        "amount": total_amount,
        "hourlyRate": hourly_rate,
        "vatAmount": vat_amount,
        "vatPercent": vat_percent,
        "vatIncluded": vat_included,
        "includeTimesheet": include_timesheet,
        "approval_status": aggregated_status,
        "approvalNotes": _deduplicate_notes(collected_notes),
        "metadata": aggregated_metadata,
    })

    return aggregated_document


def _transition_document_v2_approval(
    session: Session,
    document: orm_models.DocumentV2ORM,
    action: DocumentApprovalAction,
    *,
    user_role: str,
    user_roles: Collection[str] | None = None,
    user_identifier: str,
    user_id: str,
    note: str | None,
) -> DocumentRecord:
    meta = document.meta or {}
    if not isinstance(meta, dict):
        meta = {}

    approval = meta.get("approval")
    if not isinstance(approval, dict):
        approval = {}

    current_status = _normalize_v2_status(approval.get("status"))
    now = datetime.utcnow()
    extra_roles_iter = (
        user_roles
        if isinstance(user_roles, Iterable) and not isinstance(user_roles, (str, bytes))
        else None
    )
    effective_roles = _collect_user_roles(user_role, extra_roles_iter)
    is_admin = bool(effective_roles & {"admin", "accountant"})

    timeline = approval.get("timeline")
    if not isinstance(timeline, list):
        timeline = []

    performer_assignee = approval.get("performer_assignee") if isinstance(approval, dict) else None
    manager_assignee = approval.get("manager_assignee") if isinstance(approval, dict) else None

    performer_assignee_id = None
    manager_assignee_id = None

    if isinstance(performer_assignee, dict):
        raw_id = performer_assignee.get("id")
        if isinstance(raw_id, str):
            performer_assignee_id = raw_id

    if isinstance(manager_assignee, dict):
        raw_id = manager_assignee.get("id")
        if isinstance(raw_id, str):
            manager_assignee_id = raw_id

    def append_history(status: str, message: str | None) -> None:
        entry = {
            "status": status,
            "timestamp": now.isoformat(timespec="seconds") + "Z",
            "author": user_identifier,
            "role": user_role,
        }
        if message and message.strip():
            entry["message"] = message.strip()
        timeline.append(entry)

    if action is DocumentApprovalAction.submit:
        allowed_statuses = {"draft", "rejected_performer", "rejected_manager"}
        if current_status not in allowed_statuses and not is_admin:
            raise ValueError("Документ уже отправлен на согласование")
        if not (effective_roles & {"admin", "accountant", "manager"}):
            raise ValueError("Недостаточно прав для отправки на согласование")

        target_status = "pending_performer"
        if current_status == "rejected_manager":
            target_status = "pending_manager"

        approval["submitted_at"] = now.isoformat()
        approval.pop("finalized_at", None)
        approval.pop("finalized_by", None)

        if target_status == "pending_manager":
            if manager_assignee_id is None:
                raise ValueError("Назначьте менеджера перед отправкой на согласование")
            approval["status"] = "pending_manager"
            approval.pop("manager_approved_at", None)
            approval.pop("manager_approved_by", None)
            append_history("pending_manager", note)
        else:
            if not performer_assignee_id:
                raise ValueError("Назначьте исполнителя перед отправкой на согласование")
            approval["status"] = "pending_performer"
            approval.pop("performer_approved_at", None)
            approval.pop("performer_approved_by", None)
            approval.pop("manager_approved_at", None)
            approval.pop("manager_approved_by", None)
            append_history("pending_performer", note)
    elif action is DocumentApprovalAction.performer_approve:
        if current_status != "pending_performer" and not is_admin:
            raise ValueError("Документ не ожидает подтверждения исполнителем")
        if not is_admin and (not performer_assignee_id or performer_assignee_id != user_id):
            raise ValueError("Документ назначен другому исполнителю")
        approval["status"] = "pending_manager"
        approval["performer_approved_at"] = now.isoformat()
        approval["performer_approved_by"] = user_identifier
        append_history("pending_manager", note)
    elif action is DocumentApprovalAction.performer_reject:
        if current_status != "pending_performer" and not is_admin:
            raise ValueError("Документ не ожидает подтверждения исполнителем")
        if not is_admin and (not performer_assignee_id or performer_assignee_id != user_id):
            raise ValueError("Документ назначен другому исполнителю")
        if not note or not note.strip():
            raise ValueError("Укажите комментарий при отклонении")
        approval["status"] = "rejected_performer"
        approval.pop("performer_approved_at", None)
        approval.pop("performer_approved_by", None)
        approval.pop("manager_approved_at", None)
        approval.pop("manager_approved_by", None)
        append_history("rejected_performer", note)
    elif action is DocumentApprovalAction.manager_approve:
        if current_status != "pending_manager" and not is_admin:
            raise ValueError("Документ не ожидает согласования менеджера")
        if manager_assignee_id:
            if not is_admin and manager_assignee_id != user_id:
                raise ValueError("Документ назначен другому менеджеру")
        elif not is_admin:
            raise ValueError("Назначьте менеджера перед согласованием")
        approval["status"] = "manager_approved"
        approval["manager_approved_at"] = now.isoformat()
        approval["manager_approved_by"] = user_identifier
        append_history("manager_approved", note)
    elif action is DocumentApprovalAction.manager_reject:
        if current_status != "pending_manager" and not is_admin:
            raise ValueError("Документ не ожидает согласования менеджера")
        if manager_assignee_id:
            if not is_admin and manager_assignee_id != user_id:
                raise ValueError("Документ назначен другому менеджеру")
        elif not is_admin:
            raise ValueError("Назначьте менеджера перед отклонением")
        if not note or not note.strip():
            raise ValueError("Укажите комментарий при отклонении")
        approval["status"] = "rejected_manager"
        approval.pop("manager_approved_at", None)
        approval.pop("manager_approved_by", None)
        approval.pop("finalized_at", None)
        approval.pop("finalized_by", None)
        append_history("rejected_manager", note)
    elif action is DocumentApprovalAction.finalize:
        if current_status != "manager_approved" and not is_admin:
            raise ValueError("Документ ещё не согласован менеджером")
        if not (effective_roles & {"admin", "accountant", "manager"}):
            raise ValueError("Недостаточно прав для завершения документа")
        approval["status"] = "final"
        approval["finalized_at"] = now.isoformat()
        approval["finalized_by"] = user_identifier
        append_history("final", note)
    else:  # pragma: no cover - defensive
        raise ValueError("Неизвестное действие согласования")

    approval["timeline"] = timeline
    meta["approval"] = approval
    document.meta = dict(meta)
    flag_modified(document, "meta")
    session.flush()
    return _map_closing_document(session, document)



def _calculate_hourly_rate(rate_type: str, base_rate: float, hourly_rate: float, norm_hours: float | None) -> float:
    if rate_type == "month":
        hours = norm_hours or 168.0
        if hours <= 0:
            hours = 168.0
        return round(base_rate / hours, 2)
    return round(hourly_rate if hourly_rate > 0 else base_rate, 2)


def _build_task_snapshots(tasks: Iterable[orm_models.TaskORM], hourly_rate: float) -> list[dict]:
    snapshots: list[dict] = []
    for task in tasks:
        spent_seconds = float(getattr(task, "spent_seconds", 0.0) or 0.0)
        billed_seconds = float(getattr(task, "billed_seconds", 0.0) or 0.0)
        remaining_seconds = max(spent_seconds - billed_seconds, 0.0)
        if remaining_seconds <= 1e-6:
            continue
        hours = round(remaining_seconds / 3600, 2)
        snapshots.append(
            {
                "id": task.id,
                "key": _issue_key(task.id),
                "title": task.summary,
                "status": task.status,
                "hours": hours,
                "billable": task.billable,
                "forceIncluded": task.force_included,
                "projectKey": task.project_key,
                "projectName": task.project_name,
                "hourlyRate": hourly_rate,
                "amount": round(hours * hourly_rate, 2),
                "assigneeDisplayName": task.assignee_display_name,
                "assigneeEmail": task.assignee_email,
                "assigneeAccountId": task.assignee_account_id,
            }
        )
    return snapshots


def _write_document_files(
    record: orm_models.DocumentRecordORM,
    payload: DocumentCreateRequest,
    work_package: orm_models.WorkPackageORM,
    rendered_content: str,
) -> list[dict]:
    files: list[dict] = []
    doc_type = DOCUMENT_TYPE_MAP.get(payload.documentType, "package")
    main_file_id = f"{record.id}-main"
    requested_format = payload.format.lower()

    if requested_format not in {"docx", "doc"}:
        raise ValueError("Формат пока не поддерживается. Выберите DOCX")

    extension = "docx"
    main_path = DOCUMENTS_DIR / f"{main_file_id}.{extension}"
    _create_docx_from_text(rendered_content, main_path)
    files.append(
        {
            "id": main_file_id,
            "label": payload.documentType,
            "type": doc_type,
            "format": "docx",
            "status": "Готов",
            "url": f"/documents/{record.id}/files/{main_file_id}",
        }
    )

    if payload.includeTimesheet:
        timesheet_id = f"{record.id}-timesheet"
        timesheet_path = DOCUMENTS_DIR / f"{timesheet_id}.xlsx"
        header = "Key\tHours\tAmount\n"
        rows = [f"{snapshot['key']}\t{snapshot['hours']}\t{snapshot['amount']}" for snapshot in work_package.task_snapshots]
        timesheet_path.write_text(header + "\n".join(rows), encoding="utf-8")
        files.append(
            {
                "id": timesheet_id,
                "label": "Таймшит",
                "type": "timesheet",
                "format": "xlsx",
                "status": "Готов",
                "url": f"/documents/{record.id}/files/{timesheet_id}",
            }
        )

    return files


def generate_document(session: Session, payload: DocumentCreateRequest) -> DocumentCreateResponse:
    work_package: orm_models.WorkPackageORM | None = None

    if payload.workPackageId:
        work_package = session.get(orm_models.WorkPackageORM, payload.workPackageId)
        if not work_package:
            raise ValueError("Work package not found")
        tasks = session.execute(
            select(orm_models.TaskORM).where(orm_models.TaskORM.work_package_id == work_package.id)
        ).scalars().all()
        task_snapshots = work_package.task_snapshots or []
        metadata = _compose_metadata(
            existing=work_package.metadata_json if isinstance(work_package.metadata_json, dict) else {},
            payload=payload,
            include_timesheet=payload.includeTimesheet,
            document_type=payload.documentType,
            period=work_package.period,
            project_key=work_package.project_key,
        )
        work_package.metadata_json = metadata
        work_package.task_snapshots = task_snapshots
        work_package.include_timesheet = payload.includeTimesheet
        performer_type = payload.performerType or work_package.performer_type or "individual"
        vat_included = payload.vatIncluded if payload.vatIncluded is not None else work_package.vat_included
        vat_percent = (
            float(payload.vatPercent)
            if payload.vatPercent is not None
            else float(work_package.vat_percent or 0.0)
        )
        vat_amount = work_package.vat_amount or 0.0
    else:
        tasks = session.execute(
            select(orm_models.TaskORM).where(orm_models.TaskORM.id.in_(payload.taskIds))
        ).scalars().all()
        if not tasks:
            raise ValueError("Выберите задачи для формирования документов")
        project_keys = {task.project_key for task in tasks}
        if len(project_keys) != 1:
            raise ValueError("Выберите задачи одного проекта")
        contract = session.get(orm_models.ContractORM, payload.contractId)
        if not contract:
            raise ValueError("Не найден контракт")
        if not contract.client_id or not contract.contractor_id:
            raise ValueError("Заполните реквизиты в карточке контракта")

        rate_type = payload.rateType or contract.rate_type or "hour"
        base_rate = float(payload.baseRate or contract.rate or 0.0)
        hourly_rate = _calculate_hourly_rate(rate_type, base_rate, payload.hourlyRate, payload.normHours)
        task_snapshots = _build_task_snapshots(tasks, hourly_rate)
        if not task_snapshots:
            raise ValueError("Нет новых часов по выбранным задачам")
        total_hours = round(sum(snapshot["hours"] for snapshot in task_snapshots), 2)
        total_amount = round(sum(snapshot["amount"] for snapshot in task_snapshots), 2)
        metadata = _compose_metadata(
            existing=None,
            payload=payload,
            include_timesheet=payload.includeTimesheet,
            document_type=payload.documentType,
            period=payload.period,
            project_key=tasks[0].project_key,
        )
        performer_type = payload.performerType or "individual"
        vat_included = bool(payload.vatIncluded)
        vat_percent = float(payload.vatPercent or (20.0 if vat_included else 0.0)) if vat_included else 0.0
        vat_amount = round(total_amount - (total_amount / (1 + vat_percent / 100)), 2) if vat_included and vat_percent else 0.0

        work_package = orm_models.WorkPackageORM(
            period=payload.period,
            project_key=tasks[0].project_key,
            project_name=tasks[0].project_name,
            contract_id=payload.contractId,
            client_id=contract.client_id,
            contractor_id=contract.contractor_id,
            total_hours=total_hours,
            total_amount=total_amount,
            hourly_rate=hourly_rate,
            base_rate=base_rate,
            rate_type=rate_type,
            include_timesheet=payload.includeTimesheet,
            currency="RUB",
            metadata_json=metadata,
            task_snapshots=task_snapshots,
            performer_type=performer_type,
            vat_included=vat_included,
            vat_percent=vat_percent,
            vat_amount=vat_amount,
        )
        session.add(work_package)
        session.flush()

        increments = {snapshot["id"]: snapshot["hours"] * 3600 for snapshot in task_snapshots}
        for task in tasks:
            increment = increments.get(task.id)
            if not increment or increment <= 1e-6:
                continue
            task.work_package_id = work_package.id
            task.force_included = False
            current_billed = float(getattr(task, "billed_seconds", 0.0) or 0.0)
            task.billed_seconds = current_billed + increment

    if work_package:
        if payload.performerType:
            work_package.performer_type = payload.performerType
        if payload.vatIncluded is not None:
            work_package.vat_included = payload.vatIncluded
        if payload.vatPercent is not None:
            work_package.vat_percent = float(payload.vatPercent)
        work_package.vat_amount = (
            round(
                work_package.total_amount
                - work_package.total_amount / (1 + work_package.vat_percent / 100),
                2,
            )
            if work_package.vat_included and work_package.vat_percent
            else 0.0
        )
        performer_type = work_package.performer_type
        vat_included = work_package.vat_included
        vat_percent = work_package.vat_percent
        vat_amount = work_package.vat_amount
    else:
        performer_type = payload.performerType or "individual"
        vat_included = bool(payload.vatIncluded)
        vat_percent = float(payload.vatPercent or 0.0) if vat_included else 0.0
        vat_amount = 0.0

    # Ensure we have metadata & snapshots available
    if not work_package:
        raise ValueError("Не удалось подготовить пакет работ")

    metadata_for_record = dict(work_package.metadata_json or {})
    metadata_for_record["variables"] = payload.variables or {}

    record = orm_models.DocumentRecordORM(
        period=work_package.period,
        type=payload.documentType,
        format=payload.format,
        status="Готов",
        include_timesheet=payload.includeTimesheet,
        files=[],
        work_package_id=work_package.id,
        project_key=work_package.project_key,
        project_name=work_package.project_name,
        contract_id=work_package.contract_id,
        client_id=work_package.client_id,
        contractor_id=work_package.contractor_id,
        total_hours=work_package.total_hours,
        total_amount=work_package.total_amount,
        hourly_rate=work_package.hourly_rate,
        base_rate=work_package.base_rate,
        rate_type=work_package.rate_type,
        task_count=len(work_package.task_snapshots or []),
        metadata_json=metadata_for_record,
        template_id=payload.templateId,
        performer_type=performer_type,
        vat_included=vat_included,
        vat_percent=vat_percent,
        vat_amount=vat_amount,
    )
    session.add(record)
    session.flush()

    contract_obj = session.get(orm_models.ContractORM, work_package.contract_id)
    client = session.get(orm_models.LegalEntityORM, work_package.client_id)
    contractor = session.get(orm_models.IndividualORM, work_package.contractor_id)

    _hydrate_document_assignees(session, record)

    context = _build_context(
        work_package=work_package,
        contract=contract_obj,
        client=client,
        contractor=contractor,
        tasks=tasks,
        payload=payload,
        vat_amount=vat_amount,
    )
    if payload.variables:
        context.update({key: str(value) for key, value in payload.variables.items()})

    template_content = None
    if payload.templateId:
        template = get_template(payload.templateId)
        if not template:
            raise ValueError("Не найден шаблон документа")
        template_content = template.content

    # === GPT / AUTO mode: always prepare narrative text from real tasks and prefer it when template asks ===
    template_content = template_content  # no-op for clarity
    auto_wants_gpt = False
    if template_content:
        if ("${gptBody}" in template_content) or ("${table2}" in template_content) or ("${tableTasks}" in template_content):
            auto_wants_gpt = True

    # Source items for narrative: prefer real TaskORMs (have descriptions), otherwise snapshots
    if tasks:
        source_items = list(tasks)
    else:
        source_items = list(work_package.task_snapshots or [])

    task_bullets = _prepare_task_bullets(source_items) if source_items else []

    narrative_html = ""
    if source_items:
        try:
            narrative_html = _generate_gpt_act_text(source_items, payload)
        except ServiceError as exc:
            logger.debug("GPT narrative unavailable: %s", exc)
            lang = "ru"
            gpt_options = getattr(payload, "gptOptions", None)
            if gpt_options and getattr(gpt_options, "language", None):
                lang = gpt_options.language
            narrative_html = _build_plain_act_html(task_bullets, lang=lang)

    # --- Diagnostics about descriptions availability ---
    total_items = len(source_items)
    described = sum(1 for item in source_items if _has_description(item)) if source_items else 0

    # Store diagnostics in metadata for visibility
    meta_diag = work_package.metadata_json if isinstance(work_package.metadata_json, dict) else {}
    meta_diag.setdefault("diagnostics", {})
    meta_diag["diagnostics"]["tasks_total"] = total_items
    meta_diag["diagnostics"]["tasks_with_description"] = described
    work_package.metadata_json = meta_diag

    # Optional guard: if caller requires real descriptions, fail early
    require_desc = False
    if payload.variables and isinstance(payload.variables, dict):
        require_desc = _boolish(payload.variables.get("requireDescriptions"))
    if require_desc and described == 0:
        raise ValueError("В выбранных задачах нет ни одного непустого описания (description). Проверьте, что из JIRA подгружается поле 'description'.")

    # Put narrative to context unconditionally — templates that don't use ${gptBody} will ignore it
    if narrative_html:
        context["gptBody"] = narrative_html
    else:
        # ensure key exists so placeholder won't render as empty if template expects it
        context.setdefault("gptBody", "")

    if narrative_html and not context.get("bodygpt"):
        context["bodygpt"] = narrative_html

    assignment_sections = _generate_service_assignment_sections(
        bullets=task_bullets,
        context=context,
        payload=payload,
    )
    for key, value in assignment_sections.items():
        if not context.get(key):
            context[key] = value

    template_export = {
        key: context[key]
        for key in TEMPLATE_VARIABLE_EXPORT_KEYS
        if context.get(key)
    }
    if template_export:
        metadata_container = record.metadata_json if isinstance(record.metadata_json, dict) else {}
        existing_snapshot = metadata_container.get("template_variables") if isinstance(metadata_container, dict) else {}
        if not isinstance(existing_snapshot, dict):
            existing_snapshot = {}
        for key, value in template_export.items():
            existing_snapshot.setdefault(key, value)
        metadata_container = dict(metadata_container or {})
        metadata_container["template_variables"] = existing_snapshot
        record.metadata_json = metadata_container
        metadata_for_record["template_variables"] = existing_snapshot

    # Decide whether to replace legacy table placeholders with the narrative
    doc_type_label = (payload.documentType or "").lower()
    allow_table_replace = any(keyword in doc_type_label for keyword in {"служеб", "задани"})
    replace_tables = allow_table_replace and bool(
        (getattr(payload, "gptOptions", None) and payload.gptOptions and getattr(payload.gptOptions, "enabled", False))
        or auto_wants_gpt
    )
    if replace_tables and narrative_html:
        # Backward compatibility: substitute narrative for table placeholders
        context["tableTasks"] = narrative_html
        context["table1"] = narrative_html
        context["table2"] = narrative_html

        # Tag work package for traceability
        meta = work_package.metadata_json if isinstance(work_package.metadata_json, dict) else {}
        tags = [*(meta.get("tags") or [])]
        if "gpt" not in tags:
            tags.append("gpt")
        meta["tags"] = list(dict.fromkeys([t for t in tags if isinstance(t, str) and t.strip()]))
        work_package.metadata_json = meta

    rendered_content = (
        _render_template_content(template_content, context)
        if template_content
        else _build_default_document(context)
    )

    record.files = _write_document_files(record, payload, work_package, rendered_content)
    session.flush()

    return DocumentCreateResponse(
        record=_map_document(record),
        workPackage=_map_work_package(work_package),
    )


def list_documents(session: Session, current_user: UserPublic | None = None) -> list[DocumentRecord]:
    legacy_records = (
        session.execute(
            select(orm_models.DocumentRecordORM)
            .options(selectinload(orm_models.DocumentRecordORM.workspace))
            .order_by(orm_models.DocumentRecordORM.created_at.desc())
        )
        .scalars()
        .all()
    )

    active_workspace_id = session.info.get("workspace_id")
    if active_workspace_id:
        filtered_records: list[orm_models.DocumentRecordORM] = []
        for record in legacy_records:
            if record.workspace_id == active_workspace_id:
                filtered_records.append(record)
                continue
            if (
                getattr(record, "shared_with_parent", False)
                and record.shared_parent_id == active_workspace_id
            ):
                filtered_records.append(record)
        legacy_records = filtered_records

    for record in legacy_records:
        _hydrate_document_assignees(session, record)
    legacy = [_map_document(record) for record in legacy_records]

    closing_records = (
        session.execute(
            select(orm_models.DocumentV2ORM)
            .options(
                selectinload(orm_models.DocumentV2ORM.package),
                selectinload(orm_models.DocumentV2ORM.contract).selectinload(orm_models.ContractV2ORM.company),
                selectinload(orm_models.DocumentV2ORM.contract).selectinload(orm_models.ContractV2ORM.performer),
                selectinload(orm_models.DocumentV2ORM.performer),
                selectinload(orm_models.DocumentV2ORM.timesheet),
            )
            .order_by(orm_models.DocumentV2ORM.created_at.desc())
        )
        .scalars()
        .all()
    )

    closing_group_map: dict[tuple, list[orm_models.DocumentV2ORM]] = defaultdict(list)
    for record in closing_records:
        key = (
            record.package_id or record.id,
            record.contract_id,
            record.performer_id or 0,
            record.period_start,
            record.period_end,
        )
        closing_group_map[key].append(record)

    closing = [_merge_closing_documents(session, records) for records in closing_group_map.values()]

    if active_workspace_id:
        closing = [
            item
            for item in closing
            if item.workspaceId == active_workspace_id
            or (item.sharedWithParent and item.sharedParentId == active_workspace_id)
        ]

    closing.sort(key=lambda item: item.createdAt, reverse=True)

    combined = legacy + closing
    combined.sort(key=lambda item: item.createdAt, reverse=True)

    if current_user is not None:
        primary_role = getattr(current_user.role, "value", current_user.role)
        extra_roles = getattr(current_user, "roles", None)
        extra_iterable = (
            extra_roles
            if isinstance(extra_roles, Iterable) and not isinstance(extra_roles, (str, bytes))
            else None
        )
        effective_roles = _collect_user_roles(primary_role, extra_iterable)
        user_id = getattr(current_user, "id", None)
        is_admin = bool(effective_roles & {"admin", "accountant"})

        if user_id and not is_admin:
            include_performer = "performer" in effective_roles
            include_manager = "manager" in effective_roles
            if include_performer or include_manager:
                combined = [
                    record
                    for record in combined
                    if (
                        (include_performer and record.performerAssignee is not None and record.performerAssignee.id == user_id)
                        or (include_manager and record.managerAssignee is not None and record.managerAssignee.id == user_id)
                    )
                ]
            else:
                combined = []

    return combined


def _parse_package_v2_document_id(document_id: str) -> Optional[int]:
    if not document_id:
        return None
    if document_id.startswith("package-v2-"):
        suffix = document_id.removeprefix("package-v2-")
        try:
            return int(suffix)
        except ValueError:
            return None
    try:
        return int(document_id)
    except ValueError:
        return None


def list_work_packages(session: Session) -> list[WorkPackage]:
    legacy_records = (
        session.execute(
            select(orm_models.WorkPackageORM).order_by(orm_models.WorkPackageORM.created_at.desc())
        )
        .scalars()
        .all()
    )
    legacy_packages = [_map_work_package(pkg) for pkg in legacy_records]

    closing_records = (
        session.execute(
            select(orm_models.DocumentV2ORM)
            .options(
                selectinload(orm_models.DocumentV2ORM.package),
                selectinload(orm_models.DocumentV2ORM.contract).selectinload(orm_models.ContractV2ORM.company),
                selectinload(orm_models.DocumentV2ORM.contract).selectinload(orm_models.ContractV2ORM.performer),
                selectinload(orm_models.DocumentV2ORM.performer),
                selectinload(orm_models.DocumentV2ORM.timesheet),
            )
            .order_by(orm_models.DocumentV2ORM.created_at.desc())
        )
        .scalars()
        .all()
    )

    grouped: dict[int, dict[str, orm_models.DocumentV2ORM]] = defaultdict(dict)
    for record in closing_records:
        if record.package_id is None:
            continue

        performer_key = (
            str(record.performer_id)
            if record.performer_id is not None
            else f"doc-{record.id}"
        )
        package_bucket = grouped[record.package_id]
        existing = package_bucket.get(performer_key)
        if existing is None:
            package_bucket[performer_key] = record
            continue

        candidate_priority = _doc_v2_priority(record)
        existing_priority = _doc_v2_priority(existing)
        if candidate_priority < existing_priority:
            package_bucket[performer_key] = record
            continue
        if candidate_priority == existing_priority:
            existing_updated = existing.updated_at or existing.created_at or datetime.min
            candidate_updated = record.updated_at or record.created_at or datetime.min
            if candidate_updated > existing_updated:
                package_bucket[performer_key] = record

    v2_packages: list[WorkPackage] = []
    for package_id, performers in grouped.items():
        sorted_docs = sorted(
            performers.items(),
            key=lambda item: (item[0].startswith("doc-"), item[0]),
        )
        for index, (_, document) in enumerate(sorted_docs, start=1):
            try:
                v2_packages.append(_map_v2_document_to_work_package(session, document, index))
            except Exception as exc:  # pragma: no cover - best effort for legacy data
                logger.warning(
                    "Unable to map closing package document %s: %s",
                    getattr(document, "id", "unknown"),
                    exc,
                )

    combined = legacy_packages + v2_packages
    combined.sort(key=lambda pkg: pkg.createdAt, reverse=True)
    return combined


def get_document(session: Session, document_id: str) -> DocumentRecord | None:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if record:
        _hydrate_document_assignees(session, record)
        return _map_document(record)

    parsed_id = _parse_package_v2_document_id(document_id)
    if parsed_id is None:
        return None

    doc_v2 = session.get(orm_models.DocumentV2ORM, parsed_id)
    if not doc_v2:
        return None

    return _map_closing_document(session, doc_v2)


def resolve_document_file(session: Session, document_id: str, file_id: str) -> Tuple[Path, str]:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if not record:
        raise FileNotFoundError
    for meta in record.files or []:
        if meta.get("id") == file_id:
            extension = meta.get("format", "txt").lower()
            path = DOCUMENTS_DIR / f"{file_id}.{extension}"
            if path.exists():
                filename = f"{file_id}.{extension}"
                return path, filename
            break
    raise FileNotFoundError


def _remove_document_files(file_metadata: Iterable[dict]) -> None:
    for meta in file_metadata:
        file_id = meta.get("id")
        fmt = (meta.get("format") or "").lower().strip()
        if not file_id or not fmt:
            continue
        candidate = DOCUMENTS_DIR / f"{file_id}.{fmt}"
        try:
            if candidate.exists():
                candidate.unlink()
        except OSError:
            # Best-effort cleanup; ignore filesystem errors.
            pass


def _remove_document_v2_files(paths: Iterable[str]) -> None:
    for raw_path in paths:
        if not raw_path:
            continue
        try:
            candidate = Path(raw_path)
            if not candidate.is_absolute():
                candidate = (DOCUMENTS_DIR / candidate).resolve()
            if candidate.exists():
                candidate.unlink()
        except OSError as exc:
            logger.warning("Unable to remove document file %s: %s", raw_path, exc)


def delete_document(session: Session, document_id: str) -> bool:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if not record:
        parsed_id = _parse_package_v2_document_id(document_id)
        if parsed_id is None:
            return False

        document_v2 = session.get(orm_models.DocumentV2ORM, parsed_id)
        if not document_v2:
            return False

        package = document_v2.package
        documents_to_remove = list(package.documents) if package else [document_v2]

        artifact_paths: List[str] = []
        work_package_keys: set[str] = set()

        for doc in documents_to_remove:
            if doc.file_path:
                artifact_paths.append(doc.file_path)
            for version in doc.versions:
                if version.file_path:
                    artifact_paths.append(version.file_path)

            key = _v2_work_package_key_from_ids(doc.package_id, doc.performer_id)
            if key:
                work_package_keys.add(key)

        for key in work_package_keys:
            tasks = (
                session.execute(
                    select(orm_models.TaskORM).where(orm_models.TaskORM.work_package_id == key)
                )
                .scalars()
                .all()
            )
            for task in tasks:
                task.work_package_id = None
                task.force_included = False

        if package:
            session.delete(package)
        else:
            session.delete(document_v2)

        session.flush()

        _remove_document_v2_files(artifact_paths)

        return True

    file_metadata = list(record.files or [])
    work_package = record.work_package
    work_package_id = record.work_package_id

    if work_package_id:
        try:
            release_work_package_tasks(session, work_package_id)
        except ValueError:
            # Work package might already be deleted; ignore.
            pass

    remove_work_package = bool(work_package and len(work_package.documents) <= 1)

    session.delete(record)

    if remove_work_package and work_package is not None:
        session.delete(work_package)

    session.flush()

    _remove_document_files(file_metadata)

    return True


def transition_document_approval(
    session: Session,
    document_id: str,
    action: DocumentApprovalAction,
    *,
    user_role: str,
    user_roles: Collection[str] | None = None,
    user_identifier: str,
    user_id: str,
    note: str | None = None,
) -> DocumentRecord:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if record:
        _hydrate_document_assignees(session, record)
        current_status = record.approval_status or "draft"
        now = datetime.utcnow()
        notes = [
            DocumentApprovalNote(**item) if not isinstance(item, DocumentApprovalNote) else item
            for item in (record.approval_notes or [])
        ]

        extra_roles_iter = (
            user_roles
            if isinstance(user_roles, Iterable) and not isinstance(user_roles, (str, bytes))
            else None
        )
        effective_roles = _collect_user_roles(user_role, extra_roles_iter)
        is_admin = bool(effective_roles & {"admin", "accountant"})

        def append_note(status: str, message: str | None) -> None:
            if message and message.strip():
                notes.append(
                    DocumentApprovalNote(
                        timestamp=now,
                        author=user_identifier,
                        role=user_role,
                        status=status,
                        message=message.strip(),
                    )
                )

        performer_assignee_id = record.performer_assignee_id
        manager_assignee_id = record.manager_assignee_id

        def ensure_assignee(expected_id: str | None, role_label: str) -> None:
            if not expected_id:
                raise ValueError(f"Назначьте {role_label} перед этим действием")
            if effective_roles & {"admin", "accountant"}:
                return
            if user_id != expected_id:
                raise ValueError("Документ назначен другому пользователю")

        if action is DocumentApprovalAction.submit:
            allowed_statuses = {"draft", "rejected_performer", "rejected_manager"}
            if current_status not in allowed_statuses and not is_admin:
                raise ValueError("Документ уже отправлен на согласование")
            if not (effective_roles & {"admin", "accountant", "manager"}):
                raise ValueError("Недостаточно прав для отправки на согласование")

            target_status = "pending_performer"
            if current_status == "rejected_manager":
                target_status = "pending_manager"

            record.submitted_at = now
            record.finalized_at = None
            record.finalized_by = None

            if target_status == "pending_manager":
                if manager_assignee_id is None:
                    raise ValueError("Назначьте менеджера перед отправкой на согласование")
                record.approval_status = "pending_manager"
                record.manager_approved_at = None
                record.manager_approved_by = None
                append_note("pending_manager", note)
            else:
                if not performer_assignee_id:
                    raise ValueError("Назначьте исполнителя перед отправкой на согласование")
                record.approval_status = "pending_performer"
                record.performer_approved_at = None
                record.performer_approved_by = None
                record.manager_approved_at = None
                record.manager_approved_by = None
                append_note("pending_performer", note)

        elif action is DocumentApprovalAction.performer_approve:
            if current_status != "pending_performer" and not is_admin:
                raise ValueError("Документ не ожидает подтверждения исполнителем")
            ensure_assignee(performer_assignee_id, "исполнителя")
            record.approval_status = "pending_manager"
            record.performer_approved_at = now
            record.performer_approved_by = user_identifier
            append_note("pending_manager", note)

        elif action is DocumentApprovalAction.performer_reject:
            if current_status != "pending_performer" and not is_admin:
                raise ValueError("Документ не ожидает подтверждения исполнителем")
            ensure_assignee(performer_assignee_id, "исполнителя")
            if not note or not note.strip():
                raise ValueError("Укажите комментарий при отклонении")
            record.approval_status = "rejected_performer"
            record.performer_approved_at = None
            record.performer_approved_by = None
            record.manager_approved_at = None
            record.manager_approved_by = None
            append_note("rejected_performer", note)

        elif action is DocumentApprovalAction.manager_approve:
            if current_status != "pending_manager" and not is_admin:
                raise ValueError("Документ не ожидает согласования менеджера")
            if manager_assignee_id:
                ensure_assignee(manager_assignee_id, "менеджера")
            elif not (effective_roles & {"admin", "accountant"}):
                raise ValueError("Назначьте менеджера перед согласованием")
            record.approval_status = "manager_approved"
            record.manager_approved_at = now
            record.manager_approved_by = user_identifier
            append_note("manager_approved", note)

        elif action is DocumentApprovalAction.manager_reject:
            if current_status != "pending_manager" and not is_admin:
                raise ValueError("Документ не ожидает согласования менеджера")
            if manager_assignee_id:
                ensure_assignee(manager_assignee_id, "менеджера")
            elif not (effective_roles & {"admin", "accountant"}):
                raise ValueError("Назначьте менеджера перед отклонением")
            if not note or not note.strip():
                raise ValueError("Укажите комментарий при отклонении")
            record.approval_status = "rejected_manager"
            record.manager_approved_at = None
            record.manager_approved_by = None
            record.finalized_at = None
            record.finalized_by = None
            append_note("rejected_manager", note)

        elif action is DocumentApprovalAction.finalize:
            if current_status != "manager_approved" and not is_admin:
                raise ValueError("Документ ещё не согласован менеджером")
            if not (effective_roles & {"admin", "accountant", "manager"}):
                raise ValueError("Недостаточно прав для завершения документа")
            record.approval_status = "final"
            record.finalized_at = now
            record.finalized_by = user_identifier
            append_note("final", note)
        else:  # pragma: no cover - defensive
            raise ValueError("Неизвестное действие согласования")

        record.approval_notes = [note.dict() if isinstance(note, DocumentApprovalNote) else note for note in notes]
        session.flush()
        return _map_document(record)

    parsed_id = _parse_package_v2_document_id(document_id)
    if parsed_id is None:
        raise ValueError("Документ не найден")

    document_v2 = session.get(orm_models.DocumentV2ORM, parsed_id)
    if not document_v2:
        raise ValueError("Документ не найден")

    return _transition_document_v2_approval(
        session,
        document_v2,
        action,
        user_role=user_role,
        user_roles=user_roles,
        user_identifier=user_identifier,
        user_id=user_id,
        note=note,
    )


def share_document_with_parent(
    session: Session,
    document_id: str,
    *,
    user: orm_models.UserORM,
) -> DocumentRecord:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if not record:
        raise ValueError("Документ не найден")

    active_workspace_id = session.info.get("workspace_id")
    if active_workspace_id and record.workspace_id != active_workspace_id:
        raise ValueError("Можно отправлять только документы текущего контура")

    workspace = record.workspace or session.get(orm_models.WorkspaceORM, record.workspace_id)
    if workspace is None:
        raise ValueError("Контур документа не найден")
    if not workspace.parent_id:
        raise ValueError("У этого контура нет родителя")

    status = (record.approval_status or "draft").strip()
    if status not in {"manager_approved", "final"}:
        raise ValueError("Отправить наверх можно только согласованный документ")

    record.shared_with_parent = True
    record.shared_parent_id = workspace.parent_id
    record.shared_at = datetime.utcnow()
    record.shared_by_user_id = user.id

    session.flush()
    _hydrate_document_assignees(session, record)
    return _map_document(record)


def revoke_document_share(session: Session, document_id: str) -> DocumentRecord:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if not record:
        raise ValueError("Документ не найден")

    active_workspace_id = session.info.get("workspace_id")
    if active_workspace_id and record.workspace_id != active_workspace_id:
        raise ValueError("Можно управлять только документами текущего контура")

    record.shared_with_parent = False
    record.shared_parent_id = None
    record.shared_at = None
    record.shared_by_user_id = None

    session.flush()
    _hydrate_document_assignees(session, record)
    return _map_document(record)


def update_document_assignees(
    session: Session,
    document_id: str,
    *,
    performer_id: str | None,
    manager_id: str | None,
) -> DocumentRecord:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if not record:
        parsed_id = _parse_package_v2_document_id(document_id)
        if parsed_id is None:
            raise ValueError("Документ не найден")
        document_v2 = session.get(orm_models.DocumentV2ORM, parsed_id)
        if not document_v2:
            raise ValueError("Документ не найден")
        return _update_document_v2_assignees(
            session,
            document_v2,
            performer_id=performer_id,
            manager_id=manager_id,
        )

    current_status = (record.approval_status or "draft").strip()

    def resolve_user(user_id: str | None) -> orm_models.UserORM | None:
        if user_id is None:
            return None
        user = session.get(orm_models.UserORM, user_id)
        if not user:
            raise ValueError("Пользователь не найден")
        return user

    performer_user = resolve_user(performer_id)
    manager_user = resolve_user(manager_id)

    if performer_user is not None and current_status not in {"draft", "pending_performer", "rejected_performer"}:
        raise ValueError("Исполнителя можно менять только до подтверждения")
    if performer_user is None and record.performer_assignee_id and current_status not in {"draft", "pending_performer", "rejected_performer"}:
        raise ValueError("Нельзя снять исполнителя после подтверждения")

    if manager_user is not None and current_status in {"manager_approved", "final"}:
        raise ValueError("Невозможно изменить менеджера после согласования")
    if manager_user is None and record.manager_assignee_id and current_status in {"manager_approved", "final"}:
        raise ValueError("Нельзя снять менеджера после согласования")

    record.performer_assignee = performer_user
    record.performer_assignee_id = performer_user.id if performer_user else None

    record.manager_assignee = manager_user
    record.manager_assignee_id = manager_user.id if manager_user else None

    if performer_user is None and current_status in {"pending_performer", "rejected_performer"}:
        record.approval_status = "draft"
        record.submitted_at = None

    if manager_user is None and current_status == "pending_manager":
        record.approval_status = "pending_performer"
        record.manager_approved_at = None
        record.manager_approved_by = None

    session.flush()
    return _map_document(record)


def _update_document_v2_assignees(
    session: Session,
    document: orm_models.DocumentV2ORM,
    *,
    performer_id: str | None,
    manager_id: str | None,
) -> DocumentRecord:
    def resolve_user(user_id: str | None) -> orm_models.UserORM | None:
        if user_id is None:
            return None
        user = session.get(orm_models.UserORM, user_id)
        if not user:
            raise ValueError("Пользователь не найден")
        return user

    performer_user = resolve_user(performer_id)
    manager_user = resolve_user(manager_id)

    meta = document.meta if isinstance(document.meta, dict) else {}
    if not isinstance(meta, dict):
        meta = {}
    approval = meta.get("approval")
    if not isinstance(approval, dict):
        approval = {}

    current_status = _normalize_v2_status(approval.get("status"))

    if performer_user is not None and current_status not in {"draft", "pending_performer", "rejected_performer"}:
        raise ValueError("Исполнителя можно менять только до подтверждения")
    if performer_user is None and approval.get("performer_assignee") and current_status not in {"draft", "pending_performer", "rejected_performer"}:
        raise ValueError("Нельзя снять исполнителя после подтверждения")

    if manager_user is not None and current_status in {"manager_approved", "final"}:
        raise ValueError("Невозможно изменить менеджера после согласования")
    if manager_user is None and approval.get("manager_assignee") and current_status in {"manager_approved", "final"}:
        raise ValueError("Нельзя снять менеджера после согласования")

    def to_meta(user: orm_models.UserORM | None) -> dict | None:
        if not user:
            return None
        full_name = (user.full_name or "").strip() or user.email
        return {
            "id": user.id,
            "email": user.email,
            "full_name": full_name,
        }

    if performer_user:
        approval["performer_assignee"] = to_meta(performer_user)
    else:
        approval.pop("performer_assignee", None)

    if manager_user:
        approval["manager_assignee"] = to_meta(manager_user)
    else:
        approval.pop("manager_assignee", None)

    if performer_user is None and current_status in {"pending_performer", "rejected_performer"}:
        approval["status"] = "draft"
        approval.pop("submitted_at", None)
    if manager_user is None and current_status == "pending_manager":
        approval["status"] = "pending_performer"
        approval.pop("manager_approved_at", None)
        approval.pop("manager_approved_by", None)

    meta["approval"] = approval
    document.meta = dict(meta)
    flag_modified(document, "meta")
    session.flush()
    return _map_closing_document(session, document)


def add_document_note(
    session: Session,
    document_id: str,
    *,
    user_role: str,
    user_identifier: str,
    message: str,
) -> DocumentRecord:
    record = session.get(orm_models.DocumentRecordORM, document_id)
    if record:
        body = message.strip()
        if not body:
            raise ValueError("Комментарий не может быть пустым")

        notes = [
            DocumentApprovalNote(**item) if not isinstance(item, DocumentApprovalNote) else item
            for item in (record.approval_notes or [])
        ]
        notes.append(
            DocumentApprovalNote(
                timestamp=datetime.utcnow(),
                author=user_identifier,
                role=user_role,
                status=record.approval_status or "draft",
                message=body,
            )
        )
        record.approval_notes = [note.dict() if isinstance(note, DocumentApprovalNote) else note for note in notes]
        session.flush()
        return _map_document(record)

    parsed_id = _parse_package_v2_document_id(document_id)
    if parsed_id is None:
        raise ValueError("Документ не найден")

    document_v2 = session.get(orm_models.DocumentV2ORM, parsed_id)
    if not document_v2:
        raise ValueError("Документ не найден")

    body = message.strip()
    if not body:
        raise ValueError("Комментарий не может быть пустым")

    meta = document_v2.meta or {}
    if not isinstance(meta, dict):
        meta = {}
    approval = meta.get("approval")
    if not isinstance(approval, dict):
        approval = {"status": "draft"}

    current_status = _normalize_v2_status(approval.get("status"))
    approval["status"] = current_status

    timeline = approval.get("timeline")
    if not isinstance(timeline, list):
        timeline = []

    timeline.append(
        {
            "status": current_status,
            "timestamp": datetime.utcnow().isoformat(timespec="seconds") + "Z",
            "author": user_identifier,
            "role": user_role,
            "message": body,
        }
    )

    approval["timeline"] = timeline
    meta["approval"] = approval
    document_v2.meta = meta
    session.flush()
    return _map_closing_document(session, document_v2)
