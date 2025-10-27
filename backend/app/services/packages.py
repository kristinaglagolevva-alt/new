from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime, date
from decimal import Decimal, ROUND_HALF_UP
from html import escape
from pathlib import Path
from types import SimpleNamespace
import hashlib
import json
import re
import uuid
from typing import Dict, List, Optional, Tuple

from docx import Document as DocxDocument
from sqlalchemy import func, select
from sqlalchemy.orm import Session

from .. import orm_models
from ..schemas import (
    ContractStatus,
    PackageCreateRequest,
    PackageCreateResponse,
    PackageOptions,
    PackagePreviewDocument,
    PackageGeneratedDocument,
    PackageTaskInput,
    PackageWarning,
    VATMode,
)
from .contracts import ServiceError
from .documents import (
    DOCUMENTS_DIR,
    _generate_gpt_act_text,
    _render_template_content,
    _create_docx_from_text,
    _format_currency,
    _format_hours,
    _format_number_plain,
    _amount_to_words,
    _v2_work_package_key_from_ids,
)
from ..storage import get_template, list_templates
from ..config import BASE_DIR


DECIMAL_ZERO = Decimal("0")
ROUND_TWO = Decimal("0.01")
ROUND_FOUR = Decimal("0.0001")

MONTHS_RU_GENITIVE = {
    1: "января",
    2: "февраля",
    3: "марта",
    4: "апреля",
    5: "мая",
    6: "июня",
    7: "июля",
    8: "августа",
    9: "сентября",
    10: "октября",
    11: "ноября",
    12: "декабря",
}


def _normalize_inn(inn: Optional[str]) -> Optional[str]:
    if not inn:
        return None
    digits = ''.join(ch for ch in inn if ch.isdigit())
    return digits or None


@dataclass
class ResolvedTask:
    jira_id: str
    performer: Optional[orm_models.IndividualV2ORM]
    contract: orm_models.ContractV2ORM
    hours: Decimal
    project_id: Optional[int]
    status: Optional[str]
    raw_meta: Dict[str, object]


@dataclass
class GroupPlan:
    contract: orm_models.ContractV2ORM
    performer: Optional[orm_models.IndividualV2ORM]
    performer_type: str
    project_id: Optional[int]
    vat_mode: VATMode
    currency: str
    doc_types: List[str]
    pair_id: Optional[str]
    tasks: List[ResolvedTask] = field(default_factory=list)
    hours: Decimal = DECIMAL_ZERO
    rate_hour: Decimal = DECIMAL_ZERO
    amount_wo_vat: Decimal = DECIMAL_ZERO
    vat_amount: Decimal = DECIMAL_ZERO
    amount_total: Decimal = DECIMAL_ZERO


class PackageBuilder:
    """Builds closing packages with grouped documents according to legal logic."""

    def __init__(
        self,
        session: Session,
        payload: PackageCreateRequest,
        *,
        actor_id: Optional[str],
    ) -> None:
        self.session = session
        self.payload = payload
        self.actor_id = actor_id
        self.options: PackageOptions = payload.options or PackageOptions()
        self.warnings: List[PackageWarning] = []
        self.resolved_tasks: List[ResolvedTask] = []
        self.group_plans: Dict[Tuple, GroupPlan] = {}
        self.package: Optional[orm_models.ClosingPackageORM] = None
        self.ta: Optional[orm_models.TechAssignmentORM] = None
        self.source_hash: str = ""

    # ------------------------------------------------------------------
    def execute(self) -> PackageCreateResponse:
        self._validate_period()
        self._load_tech_assignment()
        self._resolve_tasks()
        self._compute_source_hash()
        self._ensure_package()
        self._build_groups()
        preview = self._persist_documents()
        return preview

    # ------------------------------------------------------------------
    def _validate_period(self) -> None:
        if self.payload.period_start > self.payload.period_end:
            raise ServiceError(
                "validation_failed",
                "Дата начала периода должна быть не позднее даты окончания",
                {"period_start": "period_start > period_end"},
            )
        if not self.payload.tasks:
            raise ServiceError(
                "validation_failed",
                "Не выбраны задачи для формирования пакета",
                {"tasks": "empty"},
            )

    def _load_tech_assignment(self) -> None:
        if self.payload.ta_id:
            ta = self.session.get(orm_models.TechAssignmentORM, self.payload.ta_id)
            if not ta:
                raise ServiceError("not_found", "Техническое задание не найдено")
            self.ta = ta
            return

        ta_number = (
            f"AUTO-{self.payload.period_start.strftime('%Y%m%d')}"
            f"-{self.payload.period_end.strftime('%Y%m%d')}"
        )
        ta_query = (
            self.session.query(orm_models.TechAssignmentORM)
            .filter(
                orm_models.TechAssignmentORM.number == ta_number,
                orm_models.TechAssignmentORM.period_start == self.payload.period_start,
                orm_models.TechAssignmentORM.period_end == self.payload.period_end,
            )
            .order_by(orm_models.TechAssignmentORM.id.desc())
        )
        ta = ta_query.first()
        if not ta:
            ta = orm_models.TechAssignmentORM(
                number=ta_number,
                period_start=self.payload.period_start,
                period_end=self.payload.period_end,
                has_ip=False,
            )
            self.session.add(ta)
            self.session.flush()

        self.ta = ta

    def _resolve_tasks(self) -> None:
        for task_input in self.payload.tasks:
            resolved = self._resolve_task(task_input)
            self.resolved_tasks.append(resolved)

    def _resolve_task(self, task_input: PackageTaskInput) -> ResolvedTask:
        performer = None
        if task_input.assignee_id:
            performer = self.session.get(orm_models.IndividualV2ORM, task_input.assignee_id)
            if not performer:
                raise ServiceError(
                    "not_found",
                    "Исполнитель не найден",
                    {"assignee_id": task_input.assignee_id, "jira_id": task_input.jira_id},
                )
        if performer is None:
            performer = self._resolve_performer_from_meta(task_input)
        contract = self._resolve_contract(task_input, performer)
        hours = self._to_decimal(task_input.hours)
        if hours <= DECIMAL_ZERO:
            raise ServiceError(
                "validation_failed",
                "У задачи должны быть положительные часы",
                {"jira_id": task_input.jira_id},
            )
        return ResolvedTask(
            jira_id=task_input.jira_id,
            performer=performer,
            contract=contract,
            hours=hours,
            project_id=task_input.project_id,
            status=task_input.status,
            raw_meta=task_input.meta or {},
        )

    def _resolve_contract(
        self,
        task_input: PackageTaskInput,
        performer: Optional[orm_models.IndividualV2ORM],
    ) -> orm_models.ContractV2ORM:
        contract: Optional[orm_models.ContractV2ORM] = None
        if task_input.contract_id:
            contract = self.session.get(orm_models.ContractV2ORM, task_input.contract_id)
            if not contract:
                raise ServiceError(
                    "no_contract_resolved",
                    "Не найден контракт по указанному идентификатору",
                    {"contract_id": task_input.contract_id, "jira_id": task_input.jira_id},
                )
        else:
            query = self.session.query(orm_models.ContractV2ORM).filter(
                orm_models.ContractV2ORM.status == ContractStatus.ACTIVE.value,
                orm_models.ContractV2ORM.valid_from <= self.payload.period_end,
                orm_models.ContractV2ORM.valid_to >= self.payload.period_start,
            )
            if performer and performer.id:
                query = query.filter(orm_models.ContractV2ORM.performer_id == performer.id)
            if task_input.company_inn:
                normalized = _normalize_inn(task_input.company_inn)
                if normalized:
                    query = query.join(orm_models.CompanyORM, orm_models.ContractV2ORM.company_id == orm_models.CompanyORM.id)
                    query = query.filter(orm_models.CompanyORM.inn == normalized)
        results = query.all()
        results = self._filter_contracts_by_meta(results, task_input.meta or {})
        if not results:
            legacy_contract_id = None
            meta = task_input.meta or {}
            legacy_raw = meta.get('legacy_contract_id')
            if isinstance(legacy_raw, str):
                legacy_contract_id = legacy_raw
            if legacy_contract_id:
                contract = self._ensure_contract_v2_from_legacy(
                    legacy_contract_id,
                    performer,
                    task_input.company_inn,
                )
                if contract:
                    return contract
            raise ServiceError(
                "no_contract_resolved",
                "Не удалось подобрать контракт для задачи",
                {"jira_id": task_input.jira_id},
            )
        if len(results) > 1 and not self.options.autopick_contract:
            raise ServiceError(
                "no_contract_resolved",
                "Найдены несколько контрактов, требуется выбор",
                {"jira_id": task_input.jira_id, "contract_ids": [item.id for item in results]},
            )

        if len(results) > 1:
            # Автовыбор: предпочтительно контракт с performer_id
            prioritized = [item for item in results if item.performer_id == (performer.id if performer else None)]
            contract = prioritized[0] if prioritized else results[0]
        else:
            contract = results[0]
        # Финальная проверка: если контракт так и не определён — отдаём явную ошибку
        if not contract:
            raise ServiceError(
                "no_contract_resolved",
                "Не удалось подобрать контракт для задачи",
                {"jira_id": task_input.jira_id},
            )
        self._ensure_contract_valid(contract, task_input)
        return contract

    def _resolve_performer_from_meta(self, task_input: PackageTaskInput) -> Optional[orm_models.IndividualV2ORM]:
        meta = task_input.meta or {}
        assignee_name = meta.get('assignee')
        email = meta.get('email')
        account_id = meta.get('account_id')
        legacy_individual_id = meta.get('legacy_individual_id')

        query = self.session.query(orm_models.IndividualV2ORM)
        individuals = query.all()
        if email and isinstance(email, str):
            normalized_email = email.strip().lower()
            if normalized_email:
                for candidate in individuals:
                    if isinstance(candidate.tax_notes, dict):
                        candidate_email = candidate.tax_notes.get("email")
                        if isinstance(candidate_email, str) and candidate_email.strip().lower() == normalized_email:
                            return candidate

        if assignee_name and isinstance(assignee_name, str):
            normalized_name = self._normalize_text(assignee_name)
            if normalized_name:
                for candidate in individuals:
                    if self._normalize_text(candidate.full_name) == normalized_name:
                        return candidate

        if account_id and isinstance(account_id, str):
            normalized_account = account_id.strip().lower()
            for candidate in individuals:
                if isinstance(candidate.tax_notes, dict):
                    candidate_account = candidate.tax_notes.get("account_id")
                    if isinstance(candidate_account, str) and candidate_account.strip().lower() == normalized_account:
                        return candidate

        legacy_id = legacy_individual_id if isinstance(legacy_individual_id, str) else None
        return self._ensure_individual_v2_from_legacy(
            assignee_name if isinstance(assignee_name, str) else None,
            email if isinstance(email, str) else None,
            account_id if isinstance(account_id, str) else None,
            legacy_id,
        )

    def _filter_contracts_by_meta(
        self,
        candidates: list[orm_models.ContractV2ORM],
        meta: dict[str, object],
    ) -> list[orm_models.ContractV2ORM]:
        if not candidates:
            return candidates

        assignee_name = meta.get('assignee')
        email = meta.get('email')
        account_id = meta.get('account_id')

        filtered: list[orm_models.ContractV2ORM] = []
        if assignee_name and isinstance(assignee_name, str):
            normalized = self._normalize_text(assignee_name)
            for contract in candidates:
                performer = contract.performer
                if performer and self._normalize_text(performer.full_name) == normalized:
                    filtered.append(contract)
            if filtered:
                candidates = filtered

        if email and isinstance(email, str):
            normalized_email = email.strip().lower()
            filtered = []
            for contract in candidates:
                performer = contract.performer
                if not performer or not isinstance(performer.tax_notes, dict):
                    continue
                performer_email = performer.tax_notes.get("email")
                if isinstance(performer_email, str) and performer_email.strip().lower() == normalized_email:
                    filtered.append(contract)
            if filtered:
                candidates = filtered

        if account_id and isinstance(account_id, str):
            normalized_account = account_id.strip().lower()
            filtered = []
            for contract in candidates:
                performer = contract.performer
                if not performer or not isinstance(performer.tax_notes, dict):
                    continue
                performer_account = performer.tax_notes.get("account_id")
                if isinstance(performer_account, str) and performer_account.strip().lower() == normalized_account:
                    filtered.append(contract)
            if filtered:
                candidates = filtered

        return candidates

    @staticmethod
    def _normalize_text(value: str | None) -> str:
        return (value or "").strip().lower()

    def _ensure_individual_v2_from_legacy(
        self,
        name: str | None,
        email: str | None,
        account_id: str | None,
        legacy_id: str | None = None,
    ) -> Optional[orm_models.IndividualV2ORM]:
        criteria = []
        query = self.session.query(orm_models.IndividualV2ORM)
        if email:
            normalized_email = email.strip().lower()
            if normalized_email:
                criteria.append(func.lower(orm_models.IndividualV2ORM.tax_notes["email"].as_string()) == normalized_email)
        if account_id:
            normalized_account = account_id.strip().lower()
            if normalized_account:
                criteria.append(func.lower(orm_models.IndividualV2ORM.tax_notes["account_id"].as_string()) == normalized_account)
        if criteria:
            candidate = query.filter(*criteria).limit(1).one_or_none()
            if candidate:
                return candidate

        legacy_query = self.session.query(orm_models.IndividualORM)
        legacy = None
        if legacy_id:
            legacy = self.session.get(orm_models.IndividualORM, legacy_id)
        if not legacy and email:
            legacy = (
                legacy_query
                .filter(func.lower(orm_models.IndividualORM.email) == email.lower())
                .limit(1)
                .one_or_none()
            )
        if not legacy and account_id:
            legacy = (
                legacy_query
                .filter(func.lower(orm_models.IndividualORM.external_id) == account_id.lower())
                .limit(1)
                .one_or_none()
            )
        if not legacy and name:
            normalized_name = self._normalize_text(name)
            if normalized_name:
                legacy = (
                    legacy_query
                    .filter(func.lower(orm_models.IndividualORM.name) == normalized_name)
                    .order_by(orm_models.IndividualORM.updated_at.desc())
                    .first()
                )

        if not legacy and not (name or email or account_id):
            return None

        full_name = legacy.name if legacy else (name or (email or "Неизвестный исполнитель"))
        tax_notes = {}
        if email:
            tax_notes["email"] = email
        if account_id:
            tax_notes["account_id"] = account_id

        performer = orm_models.IndividualV2ORM(
            full_name=full_name,
            type="gph",
            inn=legacy.inn if legacy and legacy.inn else None,
            tax_notes=tax_notes or None,
            source="legacy",
        )
        self.session.add(performer)
        self.session.flush()
        return performer

    def _ensure_company_from_legacy(
        self,
        legal_id: str | None,
        inn_hint: str | None,
    ) -> Optional[orm_models.CompanyORM]:
        if not legal_id and not inn_hint:
            return None
        inn = inn_hint
        legal = None
        if legal_id:
            legal = self.session.get(orm_models.LegalEntityORM, legal_id)
            if legal and legal.inn:
                inn = legal.inn
        if not inn:
            return None
        normalized_inn = _normalize_inn(inn)
        if not normalized_inn:
            return None

        existing = (
            self.session.query(orm_models.CompanyORM)
            .filter(orm_models.CompanyORM.inn == normalized_inn)
            .one_or_none()
        )
        if existing:
            return existing

        company = orm_models.CompanyORM(
            name=legal.name if legal else f"Компания {normalized_inn}",
            inn=normalized_inn,
            kpp=legal.kpp if legal else None,
            is_ip=False,
            default_vat_mode="no_vat",
        )
        self.session.add(company)
        self.session.flush()
        return company

    def _ensure_contract_v2_from_legacy(
        self,
        legacy_contract_id: str,
        performer: Optional[orm_models.IndividualV2ORM],
        company_inn: Optional[str],
    ) -> Optional[orm_models.ContractV2ORM]:
        contract_legacy = self.session.get(orm_models.ContractORM, legacy_contract_id)
        if not contract_legacy:
            return None

        performer_v2 = performer
        if contract_legacy.contractor_id and not performer_v2:
            legacy_performer = self.session.get(orm_models.IndividualORM, contract_legacy.contractor_id)
            performer_v2 = self._ensure_individual_v2_from_legacy(
                legacy_performer.name if legacy_performer else None,
                legacy_performer.email if legacy_performer else None,
                legacy_performer.external_id if legacy_performer else None,
                legacy_performer.id if legacy_performer else None,
            )

        company = self._ensure_company_from_legacy(contract_legacy.client_id, company_inn)

        existing = (
            self.session.query(orm_models.ContractV2ORM)
            .filter(orm_models.ContractV2ORM.meta["legacy_contract_id"].as_string() == contract_legacy.id)
            .one_or_none()
        )
        if existing:
            return existing

        party_type = "company" if company else "individual"
        party_id = str(company.id) if company else str(performer_v2.id) if performer_v2 else ""

        if not party_id:
            return None

        today = date.today()
        contract_v2 = orm_models.ContractV2ORM(
            party_type=party_type,
            party_id=party_id,
            performer_id=performer_v2.id if performer_v2 else None,
            contract_number=contract_legacy.number or f"LEGACY-{contract_legacy.id}",
            contract_date=today,
            valid_from=today,
            valid_to=date(today.year + 5, today.month, today.day) if today.month != 2 or today.day != 29 else date(today.year + 5, today.month, 28),
            vat_mode="no_vat",
            rate_type=contract_legacy.rate_type or "hour",
            rate_value=contract_legacy.rate or 0,
            currency=contract_legacy.currency or "RUB",
            act_by_projects=False,
            ip_transfer_mode="embedded",
            status="active",
            meta={
                "source": "legacy",
                "legacy_contract_id": contract_legacy.id,
            },
            company_id=company.id if company else None,
        )
        self.session.add(contract_v2)
        self.session.flush()
        return contract_v2

    def _ensure_contract_valid(
        self,
        contract: orm_models.ContractV2ORM,
        task_input: PackageTaskInput,
    ) -> None:
        if not contract:
            raise ServiceError(
                "no_contract_resolved",
                "Контракт не определён",
                {"jira_id": task_input.jira_id},
            )
        if contract.valid_to < self.payload.period_end:
            raise ServiceError(
                "contract_expired",
                "Срок действия контракта истёк",
                {"contract_id": contract.id, "jira_id": task_input.jira_id},
            )
        if contract.status != ContractStatus.ACTIVE.value:
            raise ServiceError(
                "validation_failed",
                "Контракт не находится в статусе active",
                {"contract_id": contract.id},
            )

    def _compute_source_hash(self) -> None:
        tasks_payload = []
        for item in sorted(self.resolved_tasks, key=lambda t: t.jira_id):
            tasks_payload.append(
                {
                    "jira_id": item.jira_id,
                    "contract_id": item.contract.id,
                    "performer_id": item.performer.id if item.performer else None,
                    "hours": str(item.hours),
                    "project_id": item.project_id,
                }
            )
        basis = {
            "ta_id": self.payload.ta_id,
            "period_start": self.payload.period_start.isoformat(),
            "period_end": self.payload.period_end.isoformat(),
            "tasks": tasks_payload,
            "options": {
                "include_timesheets": self.options.include_timesheets,
                "norm_hours": self.options.norm_hours,
                "include_by_projects": self.options.include_by_projects,
                "autopick_contract": self.options.autopick_contract,
            },
        }
        serialized = json.dumps(basis, ensure_ascii=False, sort_keys=True).encode("utf-8")
        self.source_hash = hashlib.sha256(serialized).hexdigest()

    def _ensure_package(self) -> None:
        assert self.ta is not None
        existing = (
            self.session.query(orm_models.ClosingPackageORM)
            .filter(
                orm_models.ClosingPackageORM.ta_id == self.ta.id,
                orm_models.ClosingPackageORM.period_start == self.payload.period_start,
                orm_models.ClosingPackageORM.period_end == self.payload.period_end,
                orm_models.ClosingPackageORM.source_hash == self.source_hash,
            )
            .one_or_none()
        )
        if existing:
            self.package = existing
            existing.updated_at = datetime.utcnow()
            return

        package_no = self._generate_package_no()
        package = orm_models.ClosingPackageORM(
            ta_id=self.ta.id,
            period_start=self.payload.period_start,
            period_end=self.payload.period_end,
            package_no=package_no,
            status="draft",
            source_hash=self.source_hash,
            meta={
                "created_by": self.actor_id,
                "task_count": len(self.resolved_tasks),
            },
        )
        self.session.add(package)
        self.session.flush()
        self.package = package

    def _generate_package_no(self) -> str:
        assert self.ta is not None
        base = self.ta.number or f"TA-{self.ta.id}"
        existing_count = (
            self.session.query(func.count(orm_models.ClosingPackageORM.id))
            .filter(orm_models.ClosingPackageORM.ta_id == self.ta.id)
            .scalar()
        ) or 0
        return f"{base}-{existing_count + 1}"

    def _build_groups(self) -> None:
        for task in self.resolved_tasks:
            plan = self._plan_for_task(task)
            key = (
                task.contract.id,
                plan.performer.id if plan.performer else None,
                plan.vat_mode.value,
                plan.currency,
                plan.project_id,
                tuple(plan.doc_types),
                plan.pair_id,
            )
            if key not in self.group_plans:
                self.group_plans[key] = plan
            self.group_plans[key].tasks.append(task)
            self.group_plans[key].hours += task.hours

        for plan in self.group_plans.values():
            self._finalize_plan(plan)

    def _plan_for_task(self, task: ResolvedTask) -> GroupPlan:
        performer_type = self._detect_performer_type(task)
        project_bucket = self._project_bucket(task)
        doc_types, pair_id = self._doc_types_for(performer_type, task.contract)
        vat_mode = VATMode(task.contract.vat_mode)
        currency = task.contract.currency
        return GroupPlan(
            contract=task.contract,
            performer=task.performer or task.contract.performer,
            performer_type=performer_type,
            project_id=project_bucket,
            vat_mode=vat_mode,
            currency=currency,
            doc_types=doc_types,
            pair_id=pair_id,
        )

    def _project_bucket(self, task: ResolvedTask) -> Optional[int]:
        mode = self.options.include_by_projects
        contract = task.contract
        if mode == "off":
            return None
        if mode == "force":
            return task.project_id
        # auto mode
        return task.project_id if contract.act_by_projects else None

    def _doc_types_for(
        self,
        performer_type: str,
        contract: orm_models.ContractV2ORM,
    ) -> Tuple[List[str], Optional[str]]:
        doc_types: List[str] = []
        pair_id: Optional[str] = None
        has_ip = bool(self.ta and self.ta.has_ip)
        ip_separate = contract.ip_transfer_mode == "separate"

        if performer_type == "employee":
            doc_types.append("SERVICE_ASSIGN")
            if has_ip and ip_separate:
                doc_types.append("APP")
                pair_id = uuid.uuid4().hex
        else:
            doc_types.append("AVR")
            if has_ip and ip_separate:
                doc_types.append("APP")
                pair_id = uuid.uuid4().hex
            if contract.vat_mode in {VATMode.VAT_10.value, VATMode.VAT_20.value}:
                doc_types.append("INVOICE")

        return doc_types, pair_id

    def _detect_performer_type(self, task: ResolvedTask) -> str:
        if task.performer and task.performer.type:
            return task.performer.type
        performer = task.contract.performer
        if performer and performer.type:
            return performer.type
        if task.contract.party_type == "company":
            return "company"
        return "gph"

    def _finalize_plan(self, plan: GroupPlan) -> None:
        monetary_doc = "AVR" if "AVR" in plan.doc_types else None
        performer_type = plan.performer_type
        rate_hour = self._resolve_rate_hour(plan.contract)
        plan.rate_hour = rate_hour
        if performer_type == "employee":
            plan.rate_hour = DECIMAL_ZERO
            plan.amount_wo_vat = DECIMAL_ZERO
            plan.vat_amount = DECIMAL_ZERO
            plan.amount_total = DECIMAL_ZERO
            return

        monetary_hours = plan.hours
        amount_wo_vat = (monetary_hours * rate_hour).quantize(ROUND_TWO, rounding=ROUND_HALF_UP)
        vat_amount = DECIMAL_ZERO
        if plan.vat_mode == VATMode.VAT_10:
            vat_amount = (amount_wo_vat * Decimal("0.10")).quantize(ROUND_TWO, rounding=ROUND_HALF_UP)
        elif plan.vat_mode == VATMode.VAT_20:
            vat_amount = (amount_wo_vat * Decimal("0.20")).quantize(ROUND_TWO, rounding=ROUND_HALF_UP)
        plan.amount_wo_vat = amount_wo_vat
        plan.vat_amount = vat_amount
        plan.amount_total = (amount_wo_vat + vat_amount).quantize(ROUND_TWO, rounding=ROUND_HALF_UP)

        if performer_type == "selfemployed":
            receipt_ok = self._self_employed_receipt_ok(plan.contract)
            if not receipt_ok:
                warning = PackageWarning(
                    type="selfemployed_check",
                    message=f"Для самозанятого требуется чек НПД ({self._counterparty_name(plan)})",
                    details={"contract_id": plan.contract.id},
                )
                if self.options.allow_selfemployed_without_receipt:
                    self.warnings.append(warning)
                else:
                    raise ServiceError(
                        "selfemployed_no_receipt",
                        "Для самозанятого не подтверждён чек НПД",
                        {"contract_id": plan.contract.id},
                    )

    def _resolve_rate_hour(self, contract: orm_models.ContractV2ORM) -> Decimal:
        rate_value = Decimal(str(contract.rate_value or 0))
        if contract.rate_type == "hour" or rate_value == DECIMAL_ZERO:
            return rate_value.quantize(ROUND_FOUR) if rate_value else DECIMAL_ZERO
        # month-based rate
        norm_hours = self.options.norm_hours
        if norm_hours is None:
            meta = contract.meta or {}
            norm_hours = meta.get("norm_hours") if isinstance(meta, dict) else None
        try:
            norm_hours_value = Decimal(str(norm_hours)) if norm_hours else Decimal("168")
        except (ValueError, TypeError):
            norm_hours_value = Decimal("168")
        if norm_hours_value <= DECIMAL_ZERO:
            norm_hours_value = Decimal("168")
        hourly = (rate_value / norm_hours_value).quantize(ROUND_FOUR, rounding=ROUND_HALF_UP)
        return hourly

    def _self_employed_receipt_ok(self, contract: orm_models.ContractV2ORM) -> bool:
        meta = contract.meta or {}
        if not isinstance(meta, dict):
            return False
        return bool(
            meta.get("npd_receipt_confirmed")
            or meta.get("npd_receipt")
            or meta.get("npd_check")
        )

    def _persist_documents(self) -> PackageCreateResponse:
        assert self.package is not None
        assert self.ta is not None
        will_create: List[PackagePreviewDocument] = []
        generated_docs: List[PackageGeneratedDocument] = []
        for plan in self.group_plans.values():
            doc_base_meta = {
                "performer_type": plan.performer_type,
                "project_id": plan.project_id,
                "pair_id": plan.pair_id,
            }
            active_workspace_id = self.session.info.get("workspace_id")
            if active_workspace_id:
                doc_base_meta.setdefault("workspace_id", active_workspace_id)
            if self.options.gpt:
                doc_base_meta["gpt_enabled"] = bool(self.options.gpt.enabled)
                doc_base_meta["gpt_options"] = self.options.gpt.dict()
            base_counterparty = self._counterparty_name(plan)
            counterparty_type, counterparty_id = self._counterparty_identity(plan)

            task_items = self._build_task_items(plan)
            task_ids = [str(item.get("id")) for item in task_items if item.get("id")]
            task_ids = [task_id for task_id in dict.fromkeys(task_ids)]
            work_package_key = _v2_work_package_key_from_ids(
                self.package.id,
                plan.performer.id if plan.performer else None,
            )

            if task_ids:
                task_records = (
                    self.session.execute(
                        select(orm_models.TaskORM).where(orm_models.TaskORM.id.in_(task_ids))
                    )
                    .scalars()
                    .all()
                )
                for record in task_records:
                    record.work_package_id = work_package_key
                    record.force_included = False

            for doc_type in plan.doc_types:
                templates_map = getattr(self.options, "templates", {}) or {}
                selected_template_id = templates_map.get(doc_type) or getattr(self.options, "template_id", None)
                version = self._next_version(
                    plan.contract.id,
                    doc_type,
                    plan.project_id,
                )
                amount_wo_vat = plan.amount_wo_vat if doc_type in {"AVR", "APP", "INVOICE"} else DECIMAL_ZERO
                vat_amount = plan.vat_amount if doc_type in {"AVR", "APP", "INVOICE"} else DECIMAL_ZERO
                amount_total = plan.amount_total if doc_type in {"AVR", "APP", "INVOICE"} else DECIMAL_ZERO
                hours = plan.hours if doc_type in {"AVR", "SERVICE_ASSIGN"} else None
                rate_hour = plan.rate_hour if doc_type in {"AVR", "SERVICE_ASSIGN"} else None

                first_task_meta = plan.tasks[0].raw_meta if plan.tasks else {}
                project_key = None
                project_name = None
                if isinstance(first_task_meta, dict):
                    project_key = first_task_meta.get("project_key") or first_task_meta.get("projectKey")
                    project_name = first_task_meta.get("project_name") or first_task_meta.get("projectName")
                legacy_contract_id = None
                if isinstance(first_task_meta, dict):
                    legacy_contract_id = first_task_meta.get("legacy_contract_id")

                document_meta = {
                    **doc_base_meta,
                    "doc_type": doc_type,
                    "tasks": [task.jira_id for task in plan.tasks],
                    "source_hash": self.source_hash,
                }
                if task_items:
                    document_meta["task_items"] = [dict(item) for item in task_items]
                if task_ids:
                    document_meta["task_ids"] = list(task_ids)
                if project_key:
                    document_meta["project_key"] = project_key
                if project_name:
                    document_meta["project_name"] = project_name
                if legacy_contract_id:
                    document_meta["legacy_contract_id"] = legacy_contract_id

                if "approval" not in document_meta:
                    document_meta["approval"] = {"status": "draft"}

                document = orm_models.DocumentV2ORM(
                    package_id=self.package.id,
                    ta_id=self.ta.id,
                    pair_id=plan.pair_id if doc_type in {"AVR", "APP"} else None,
                    doc_type=doc_type,
                    template_id=selected_template_id,
                    counterparty_type=counterparty_type,
                    counterparty_id=counterparty_id,
                    contract_id=plan.contract.id,
                    performer_id=plan.performer.id if plan.performer else None,
                    project_id=plan.project_id,
                    vat_mode=plan.vat_mode.value,
                    currency=plan.currency,
                    period_start=self.payload.period_start,
                    period_end=self.payload.period_end,
                    hours=hours,
                    rate_hour=rate_hour,
                    amount_wo_vat=amount_wo_vat,
                    vat_amount=vat_amount,
                    amount_total=amount_total,
                    version=version,
                    meta=document_meta,
                )
                self.session.add(document)
                self.session.flush()

                summary_paragraphs = self._build_summary_paragraphs(plan)
                document.meta["summary_preview"] = summary_paragraphs
                if selected_template_id:
                    document.meta["template_id"] = selected_template_id
                file_path = self._render_document_file(document, plan, doc_type, summary_paragraphs)
                document.file_path = str(file_path)

                if self.options.include_timesheets and doc_type in {"AVR", "SERVICE_ASSIGN"}:
                    timesheet_rows = [
                        {
                            "jira_id": task.jira_id,
                            "hours": float(task.hours),
                            "status": task.status,
                        }
                        for task in plan.tasks
                    ]
                    timesheet = orm_models.TimesheetORM(
                        document_id=document.id,
                        task_table=timesheet_rows,
                    )
                    self.session.add(timesheet)

                will_create.append(
                    PackagePreviewDocument(
                        doc_type=doc_type,
                        counterparty=base_counterparty,
                        contract_id=plan.contract.id,
                        amount_total=float(amount_total) if amount_total is not None else 0.0,
                        vat_mode=plan.vat_mode,
                        group_info={
                            "performer_type": plan.performer_type,
                            "project_id": plan.project_id,
                            "tasks": [task.jira_id for task in plan.tasks],
                            "pair_id": plan.pair_id,
                        },
                    )
                )

                generated_docs.append(
                    PackageGeneratedDocument(
                        id=document.id,
                        doc_type=doc_type,
                        contract_id=plan.contract.id,
                        performer_id=plan.performer.id if plan.performer else None,
                        file_path=str(file_path),
                        file_url=f"/packages/{self.package.id}/documents/{document.id}/file",
                    )
                )

        self._log_audit("package.create", self.package.id, {
            "ta_id": self.ta.id,
            "documents": len(will_create),
            "warnings": [warning.type for warning in self.warnings],
        })

        return PackageCreateResponse(
            package_id=self.package.id,
            will_create=will_create,
            warnings=self.warnings,
            documents=generated_docs,
        )

    def _next_version(self, contract_id: int, doc_type: str, project_id: Optional[int]) -> int:
        existing_version = (
            self.session.query(func.max(orm_models.DocumentV2ORM.version))
            .filter(
                orm_models.DocumentV2ORM.ta_id == self.payload.ta_id,
                orm_models.DocumentV2ORM.contract_id == contract_id,
                orm_models.DocumentV2ORM.doc_type == doc_type,
                orm_models.DocumentV2ORM.period_start == self.payload.period_start,
                orm_models.DocumentV2ORM.period_end == self.payload.period_end,
                orm_models.DocumentV2ORM.project_id.is_(project_id) if project_id is None else orm_models.DocumentV2ORM.project_id == project_id,
            )
            .scalar()
        )
        return (existing_version or 0) + 1

    def _counterparty_name(self, plan: GroupPlan) -> str:
        contract = plan.contract
        if contract.party_type == "company" and contract.company:
            return contract.company.name
        performer = plan.performer or contract.performer
        if performer:
            return performer.full_name
        return contract.contract_number

    def _counterparty_identity(self, plan: GroupPlan) -> Tuple[str, Optional[int]]:
        contract = plan.contract
        if contract.party_type == "company" and contract.company_id:
            return "company", contract.company_id
        performer = plan.performer or contract.performer
        if performer:
            return "individual", performer.id
        return contract.party_type, None

    def _document_title(self, doc_type: str) -> str:
        mapping = {
            "AVR": "Акт выполненных работ",
            "APP": "Акт передачи прав",
            "INVOICE": "Счет-фактура",
            "SERVICE_ASSIGN": "Служебное задание",
            "PAYROLL_REPORT": "Отчет по персоналу",
        }
        return mapping.get(doc_type, doc_type)

    def _build_task_items(self, plan: GroupPlan) -> List[dict[str, object]]:
        items: List[dict[str, object]] = []
        for task in plan.tasks:
            meta = task.raw_meta or {}
            billable_flag = meta.get("billable")
            billable = bool(billable_flag) if billable_flag is not None else True
            force_included_flag = meta.get("force_included") if "force_included" in meta else meta.get("forceIncluded")
            force_included = bool(force_included_flag) if force_included_flag is not None else False
            hours_value = task.hours if isinstance(task.hours, Decimal) else Decimal(str(task.hours or DECIMAL_ZERO))
            rate_hour = plan.rate_hour if plan.rate_hour else DECIMAL_ZERO
            amount_value = (hours_value * rate_hour).quantize(ROUND_TWO, rounding=ROUND_HALF_UP) if rate_hour else DECIMAL_ZERO
            project_key = meta.get("projectKey") or meta.get("project_key") or ""
            project_name = meta.get("projectName") or meta.get("project_name") or ""
            items.append(
                {
                    "id": meta.get("task_id") or task.jira_id,
                    "key": task.jira_id,
                    "summary": meta.get("summary") or "",
                    "description": meta.get("description") or "",
                    "status": meta.get("status") or task.status or "",
                    "hours": float(hours_value or DECIMAL_ZERO),
                    "assignee": meta.get("assignee") or "",
                    "projectKey": project_key,
                    "projectName": project_name,
                    "billable": billable,
                    "forceIncluded": force_included,
                    "hourlyRate": float(rate_hour or DECIMAL_ZERO),
                    "amount": float(amount_value or DECIMAL_ZERO),
                }
            )
        return items

    def _build_summary_paragraphs(self, plan: GroupPlan) -> List[str]:
        items = self._build_task_items(plan)
        if not items:
            return ["Работы по выбранным задачам отсутствуют."]

        gpt_options = getattr(self.options, "gpt", None)
        if gpt_options and getattr(gpt_options, "enabled", False):
            payload = SimpleNamespace(
                period=f"{self.payload.period_start:%d.%m.%Y} — {self.payload.period_end:%d.%m.%Y}",
                documentType="Акт",
                gptOptions=SimpleNamespace(
                    enabled=True,
                    language=getattr(gpt_options, "language", "ru"),
                    style=getattr(gpt_options, "style", "neutral"),
                    extraNotes=getattr(gpt_options, "extraNotes", None),
                ),
            )
            html = _generate_gpt_act_text(items, payload)
            paragraphs = self._html_to_paragraphs(html)
            if paragraphs:
                return paragraphs

        return self._build_default_paragraphs(items)

    @staticmethod
    def _html_to_paragraphs(html_content: str | None) -> List[str]:
        if not html_content:
            return []
        text = html_content.replace("<br/>", "\n").replace("<br>", "\n")
        text = text.replace("</p>", "\n").replace("</li>", "\n")
        text = text.replace("<li>", "- ")
        text = re.sub(r"<[^>]+>", "", text)
        paragraphs = [line.strip() for line in text.splitlines() if line.strip()]
        return paragraphs

    @staticmethod
    def _build_default_paragraphs(items: List[dict[str, object]]) -> List[str]:
        result: List[str] = []
        for item in items:
            key = item.get("key") or item.get("id") or "—"
            summary = (item.get("summary") or "").strip()
            description = (item.get("description") or "").strip()
            if description and len(description) > 200:
                description = description[:197] + "..."

            fragments: list[str] = []
            if summary:
                fragments.append(summary.rstrip('.'))
            if description and description not in summary:
                fragments.append(description.rstrip('.'))
            if not fragments:
                fragments.append("Работы выполнены согласно заданию")

            actions = "; ".join(fragments)
            if actions:
                actions = actions[0].upper() + actions[1:]
            sentence = f"Выполнена задача {key}: {actions}."
            result.append(sentence)
        return result

    @staticmethod
    def _format_russian_date(value: date) -> str:
        month_name = MONTHS_RU_GENITIVE.get(value.month, "")
        formatted = f" «{value.day:02d}» {month_name} {value.year}"
        return formatted.rstrip()

    @staticmethod
    def _short_name(full_name: str | None) -> str:
        if not full_name:
            return ""
        cleaned = str(full_name).replace("\xa0", " ").strip()
        if not cleaned:
            return ""
        parts = [part for part in cleaned.replace(".", " ").split() if part]
        if not parts:
            return ""
        surname, *rest = parts
        initials = "".join(f"{segment[0].upper()}." for segment in rest if segment)
        return f"{surname} {initials}".strip()

    def _default_template_id(self, doc_type: str) -> str | None:
        bucket_map = {
            "AVR": "act",
            "APP": "custom",
            "IPR": "custom",
            "INVOICE": "invoice",
            "SERVICE_ASSIGN": "timesheet",
        }
        bucket = bucket_map.get(doc_type)
        if not bucket:
            return None
        for template in list_templates():
            if template.type == bucket:
                return template.id
        return None

    def _build_tasks_table_rows(self, plan: GroupPlan) -> str:
        items = self._build_task_items(plan)
        rows: list[str] = []

        rate = plan.rate_hour or DECIMAL_ZERO
        currency = plan.contract.currency

        for item in items:
            key = (item.get("key") or item.get("id") or "").strip()
            summary = (item.get("summary") or "").strip()
            description = (item.get("description") or "").strip()
            title_parts = [part for part in [summary, description] if part]
            title = "; ".join(title_parts)
            if key and title:
                title = f"{key} — {title}"
            elif key:
                title = key
            if not title:
                title = key or "—"

            hours_float = float(item.get("hours") or 0)
            hours_decimal = Decimal(str(hours_float)).quantize(ROUND_TWO, rounding=ROUND_HALF_UP)
            amount = (hours_decimal * rate).quantize(ROUND_TWO, rounding=ROUND_HALF_UP) if rate else DECIMAL_ZERO

            hours_label = _format_hours(float(hours_decimal))
            rate_label = _format_currency(float(rate), currency) if rate else "—"
            amount_label = _format_currency(float(amount), currency) if amount else "—"

            rows.append(
                "<tr>"
                f"<td><p>{escape(title)}</p></td>"
                f"<td><p>{escape(hours_label)}</p></td>"
                f"<td><p>{escape(rate_label)}</p></td>"
                f"<td><p>{escape(amount_label)}</p></td>"
                "</tr>"
            )

        if not rows:
            rows.append(
                "<tr>"
                "<td><p>—</p></td>"
                "<td><p>—</p></td>"
                "<td><p>—</p></td>"
                "<td><p>—</p></td>"
                "</tr>"
            )

        return "".join(rows)


    def _render_document_file(
        self,
        document: orm_models.DocumentV2ORM,
        plan: GroupPlan,
        doc_type: str,
        paragraphs: List[str],
    ) -> Path:
        DOCUMENTS_DIR.mkdir(parents=True, exist_ok=True)
        file_path = DOCUMENTS_DIR / f"doc-v2-{document.id}.docx"

        templates_map = getattr(self.options, "templates", {}) or {}
        template_id = templates_map.get(doc_type) or getattr(self.options, "template_id", None)
        if not template_id:
            template_id = self._default_template_id(doc_type)

        if template_id:
            template = get_template(template_id)
            if template:
                normalized_paragraphs = [
                    paragraph.strip()
                    for paragraph in paragraphs
                    if paragraph and paragraph.strip()
                ]
                html_body = "\n".join(f"<p>{p}</p>" for p in normalized_paragraphs)

                contract = plan.contract
                company = contract.company if hasattr(contract, "company") else None
                performer = plan.performer or contract.performer if hasattr(contract, "performer") else None

                first_task = plan.tasks[0] if plan.tasks else None
                task_meta = first_task.raw_meta if first_task else {}

                context = {
                    "gptBody": html_body,
                    "tableTasks": html_body,
                    "table2": html_body,
                    "startPeriodDate": f"{self.payload.period_start:%d.%m.%Y}",
                    "endPeriodDate": f"{self.payload.period_end:%d.%m.%Y}",
                    "totalHours": _format_hours(float(plan.hours)) if plan.hours is not None else "0",
                    "totalAmount": _format_currency(float(plan.amount_total or DECIMAL_ZERO), plan.contract.currency),
                    "totalAmountWithoutVat": _format_currency(float(plan.amount_wo_vat or DECIMAL_ZERO), plan.contract.currency),
                    "vatAmount": _format_currency(float(plan.vat_amount or DECIMAL_ZERO), plan.contract.currency),
                    "totalAmountNumeric": _format_number_plain(float(plan.amount_total or DECIMAL_ZERO)),
                    "totalAmountWords": _amount_to_words(float(plan.amount_total or DECIMAL_ZERO)),
                    "vatAmountNumeric": _format_number_plain(float(plan.vat_amount or DECIMAL_ZERO)),
                    "vatAmountWords": _amount_to_words(float(plan.vat_amount or DECIMAL_ZERO)),
                    "projectName": task_meta.get("projectName") or "",
                    "projectKey": task_meta.get("projectKey") or "",
                    "period": f"{self.payload.period_start:%d.%m.%Y} — {self.payload.period_end:%d.%m.%Y}",
                    "docType": doc_type,
                }

                if performer:
                    context["employeeName"] = getattr(performer, "full_name", "") or ""
                    context["employeeInn"] = getattr(performer, "inn", "") or ""

                contract_meta = getattr(contract, "meta", None)
                legacy_contract = None
                legacy_client = None
                legacy_performer = None
                if isinstance(contract_meta, dict):
                    legacy_id = contract_meta.get("legacy_contract_id")
                    if legacy_id:
                        legacy_contract = self.session.get(orm_models.ContractORM, legacy_id)
                        if legacy_contract:
                            legacy_client = legacy_contract.client
                            legacy_performer = legacy_contract.contractor

                if company and not legacy_client:
                    context["companyName"] = company.name or ""
                    context["clientInn"] = company.inn or ""
                    context["clientKpp"] = company.kpp or ""

                base_contract_number = getattr(contract, "contract_number", "") or ""
                if base_contract_number:
                    context["contractNumber"] = base_contract_number
                    context.setdefault("employeeContractNumber", base_contract_number)

                contract_date = getattr(contract, "contract_date", None)
                if contract_date:
                    context["employeeContractDate"] = contract_date.strftime("%d.%m.%Y")

                if legacy_contract:
                    legacy_number = legacy_contract.number or ""
                    if legacy_number and not context.get("employeeContractNumber"):
                        context["employeeContractNumber"] = legacy_number
                    legacy_date = getattr(legacy_contract, "created_at", None)
                    if legacy_date and not context.get("employeeContractDate"):
                        context["employeeContractDate"] = legacy_date.strftime("%d.%m.%Y")

                if legacy_client:
                    if legacy_client.name:
                        context["companyName"] = legacy_client.name
                    if legacy_client.signatory:
                        context["seoFullName"] = legacy_client.signatory
                        context["seoShortName"] = self._short_name(legacy_client.signatory)
                    if legacy_client.inn:
                        context["clientInn"] = legacy_client.inn
                    if legacy_client.kpp:
                        context["clientKpp"] = legacy_client.kpp

                performer_name = (getattr(performer, "full_name", "") or "").strip()
                if legacy_performer and not performer_name:
                    performer_name = legacy_performer.name or ""

                contractor_company_name = ""
                if contract.party_type == "company" and company:
                    contractor_company_name = company.name or ""
                if not contractor_company_name:
                    contractor_company_name = performer_name

                if performer_name:
                    context.setdefault("contractorSeoFullName", performer_name)
                    context.setdefault("contractorseoShortName", self._short_name(performer_name))

                if contractor_company_name:
                    context.setdefault("contractorCompanyName", contractor_company_name)

                bullet_body = ''.join(f'<li>{escape(p)}</li>' for p in normalized_paragraphs)
                if bullet_body and not context.get("bodygpt"):
                    context["bodygpt"] = f"<ol>{bullet_body}</ol>"
                elif html_body and not context.get("bodygpt"):
                    context["bodygpt"] = html_body

                table_rows_html = self._build_tasks_table_rows(plan)
                if table_rows_html:
                    context.setdefault("table1", table_rows_html)
                    if not context.get("tableTasks"):
                        context["tableTasks"] = table_rows_html
                    if not context.get("table2"):
                        context["table2"] = table_rows_html

                if not context.get("actNumber"):
                    if base_contract_number and document.version is not None:
                        context["actNumber"] = f"{base_contract_number}-{document.version:02d}"
                    else:
                        context["actNumber"] = str(document.id)

                if not context.get("date"):
                    context["date"] = self._format_russian_date(self.payload.period_end)

                extra = getattr(self.options, "template_variables", None) or {}
                if isinstance(extra, dict):
                    for key, value in extra.items():
                        context[str(key)] = str(value)

                content = _render_template_content(template.content, context)
                _create_docx_from_text(content, file_path)
                return file_path

        # Фоллбэк: быстрый DOCX без шаблона
        base_template = (BASE_DIR.parent / "test_renderer.docx")
        docx = DocxDocument(str(base_template)) if base_template.exists() else DocxDocument()
        # Удаляем стартовый контент шаблона
        for _ in list(docx.paragraphs):
            p = docx.paragraphs[0]
            p._element.getparent().remove(p._element)
        for tbl in list(docx.tables):
            tbl._element.getparent().remove(tbl._element)
        docx.add_heading(self._document_title(doc_type), level=1)
        for paragraph in paragraphs:
            docx.add_paragraph(paragraph)
        docx.add_paragraph(
            f"Период: {self.payload.period_start:%d.%m.%Y} — {self.payload.period_end:%d.%m.%Y}"
        )
        docx.save(file_path)
        return file_path

    def _log_audit(self, action: str, entity_id: int, payload: dict) -> None:
        audit = orm_models.AuditLogORM(
            actor_id=self.actor_id,
            action=action,
            entity="closing_package",
            entity_id=str(entity_id),
            payload=payload,
            error_code=None,
        )
        self.session.add(audit)

    @staticmethod
    def _to_decimal(value: float) -> Decimal:
        return Decimal(str(value)).quantize(ROUND_TWO, rounding=ROUND_HALF_UP)


def create_package(
    session: Session,
    payload: PackageCreateRequest,
    *,
    actor_id: Optional[str],
) -> PackageCreateResponse:
    builder = PackageBuilder(session, payload, actor_id=actor_id)
    return builder.execute()


def resolve_package_document_file(
    session: Session,
    package_id: int,
    document_id: int,
) -> Tuple[Path, str]:
    document = session.get(orm_models.DocumentV2ORM, document_id)
    if not document or document.package_id != package_id:
        raise FileNotFoundError
    if not document.file_path:
        raise FileNotFoundError

    file_path = Path(document.file_path)
    if not file_path.is_absolute():
        file_path = (DOCUMENTS_DIR / file_path).resolve()
    if not file_path.exists():
        raise FileNotFoundError

    return file_path, file_path.name
